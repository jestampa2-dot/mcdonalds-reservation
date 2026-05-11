from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
SOURCE = ROOT / "Full Thesis - source copy.docx"
OUTPUT = ROOT / "Full Thesis - IPO Conceptual Framework.docx"
DIAGRAM = ROOT / "scripts" / "ipo_conceptual_framework.png"


ACCENT = RGBColor(191, 31, 36)
DARK = RGBColor(31, 31, 31)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def insert_after(anchor, text, style_name="Normal", bold=False):
    new_paragraph = anchor.insert_paragraph_before("")
    anchor._p.addprevious(new_paragraph._p)
    # The line above places the new paragraph before anchor; move it after anchor.
    anchor._p.addnext(new_paragraph._p)
    run = new_paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    run.font.color.rgb = ACCENT if bold else DARK
    new_paragraph.style = style_name
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style_name == "Normal" else WD_ALIGN_PARAGRAPH.LEFT
    new_paragraph.paragraph_format.first_line_indent = Inches(0.5) if style_name == "Normal" else None
    new_paragraph.paragraph_format.space_after = Pt(6)
    return new_paragraph


def append_after(anchor, text, style_name="Normal", bold=False):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_element = OxmlElement("w:p")
    anchor._p.addnext(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    paragraph.style = style_name
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if style_name == "Normal" else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Inches(0.5) if style_name == "Normal" else None
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    run.font.color.rgb = ACCENT if bold else DARK
    return paragraph


def create_ipo_diagram(path: Path):
    if Image is None:
        return

    width, height = 1500, 820
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    try:
        title_font = ImageFont.truetype("arialbd.ttf", 38)
        box_font = ImageFont.truetype("arialbd.ttf", 28)
        body_font = ImageFont.truetype("arial.ttf", 22)
        caption_font = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        title_font = box_font = body_font = caption_font = ImageFont.load_default()

    draw.rounded_rectangle((28, 28, width - 28, height - 28), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "IPO Conceptual Framework of the McDonald's Event Planner", fill=(31, 31, 31), font=title_font)

    columns = [
        (
            "INPUT",
            "Customer account details\nEvent type and branch\nDate, time, duration\nGuest count\nPackage, menu, add-ons\nPayment proof\nAdmin/staff updates",
            (70, 190, 430, 610),
            (255, 247, 230),
        ),
        (
            "PROCESS",
            "Authenticate users\nValidate reservation details\nCheck availability\nCalculate total amount\nStore booking records\nReview payment proof\nApprove/reject booking\nAssign crew and update status",
            (570, 160, 930, 640),
            (255, 242, 194),
        ),
        (
            "OUTPUT",
            "Booking reference\nPending/confirmed status\nCustomer dashboard\nQR or booking pass\nStaff preparation list\nAdmin reports\nOrganized reservation records",
            (1070, 190, 1430, 610),
            (255, 247, 230),
        ),
    ]

    for title, body, coords, fill in columns:
        x1, y1, x2, y2 = coords
        draw.rounded_rectangle(coords, radius=22, fill=fill, outline=(184, 184, 184), width=3)
        draw.text((x1 + 28, y1 + 28), title, fill=(191, 31, 36), font=box_font)
        draw.multiline_text((x1 + 28, y1 + 92), body, fill=(55, 55, 55), font=body_font, spacing=8)

    def arrow(start, end):
        draw.line((start, end), fill=(191, 31, 36), width=7)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 26, ey - 16), (ex - 26, ey + 16)], fill=(191, 31, 36))

    arrow((430, 400), (570, 400))
    arrow((930, 400), (1070, 400))
    draw.text((60, 735), "Figure 3.1. IPO conceptual framework showing the input, process, and output of the proposed reservation system.", fill=(80, 80, 80), font=caption_font)
    image.save(path, quality=95)


def add_image_after(anchor, image_path):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_element = OxmlElement("w:p")
    anchor._p.addnext(new_element)
    paragraph = Paragraph(new_element, anchor._parent)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = paragraph.add_run().add_picture(str(image_path), width=Inches(6.4))
    doc_pr = inline._inline.docPr
    doc_pr.set("title", "IPO Conceptual Framework Diagram")
    doc_pr.set("descr", "IPO conceptual framework showing system inputs, processes, and outputs for the McDonald's Event Planner.")
    return paragraph


def build():
    doc = Document(str(SOURCE))
    paragraphs = list(doc.paragraphs)

    start = next(i for i, p in enumerate(paragraphs) if p.text.strip() in {"3.1 Conceptual Framework", "3.1 IPO Conceptual Framework"})
    end = next(i for i, p in enumerate(paragraphs[start + 1 :], start + 1) if p.text.strip() == "4. METHODOLOGY")

    heading = paragraphs[start]
    heading.text = "3.1 IPO Conceptual Framework"
    for paragraph in paragraphs[start + 1 : end]:
        delete_paragraph(paragraph)

    create_ipo_diagram(DIAGRAM)

    # Insert in reverse order because each append_after inserts immediately after the same anchor.
    inserted = []
    inserted.append(append_after(heading, "The conceptual framework of the proposed McDonald's Event Planner is presented using the Input-Process-Output (IPO) model. This model explains how the system receives reservation-related data from customers, staff, and administrators, processes the data through the web and mobile application workflow, and produces organized booking outputs that support event reservation management. The IPO framework is appropriate for this study because the proposed system focuses on transforming manual booking information into accurate, traceable, and accessible digital records."))
    inserted.append(append_after(inserted[-1], "Input refers to the information entered into the system before a reservation can be processed. For customers, the inputs include account details, selected event type, preferred McDonald's branch, event date, start time, duration, number of guests, room option, package, menu bundle, menu items, add-ons, special notes, and uploaded payment proof. For staff and administrators, the inputs include booking status updates, crew assignment, branch information, package details, room availability, menu records, and service progress updates. These inputs represent the data needed by the system to create and manage event reservations properly."))
    inserted.append(append_after(inserted[-1], "Process refers to the system operations that convert the entered data into useful reservation records. The Laravel backend authenticates users, validates booking details, checks availability, prevents conflicting schedules, calculates reservation totals, stores records in the database, manages uploaded payment proof, and sends updated booking information to the dashboards. The Vue.js web interface and React Native mobile application allow users to interact with these processes through booking forms, dashboards, staff tools, and administrator management screens. Administrators review pending reservations, approve or reject bookings, assign crew, and maintain branch and catalog information, while staff members handle check-in and service status updates."))
    inserted.append(append_after(inserted[-1], "Output refers to the information and results generated after the system completes its processing. The outputs include booking reference numbers, pending or confirmed reservation statuses, customer dashboard records, QR or booking pass details, updated availability schedules, staff preparation lists, service status records, administrator booking reports, and organized reservation histories. These outputs help customers monitor their reservations, help staff prepare and manage events, and help administrators make better decisions using accurate and centralized booking information."))
    last_anchor = inserted[-1]
    if DIAGRAM.exists():
        image_paragraph = add_image_after(inserted[-1], DIAGRAM)
        caption = append_after(image_paragraph, "Figure 3.1. IPO Conceptual Framework of the McDonald's Event Planner")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = None
        for run in caption.runs:
            run.font.size = Pt(10)
            run.italic = True
        last_anchor = caption
    append_after(last_anchor, "Overall, the IPO conceptual framework shows that the McDonald's Event Planner starts with user and reservation inputs, processes them through system validation, availability checking, database storage, booking review, and operational updates, and ends with reliable outputs for customers, staff, and administrators. Through this flow, the proposed system addresses the problems of manual reservation handling, including double booking, unclear records, slow confirmation, and difficulty in monitoring event schedules.")

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
