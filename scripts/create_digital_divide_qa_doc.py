from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\User\mcdonalds-reservation\docs\Digital Divide Oral Defense Q&A.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="B7C1CC", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_question_block(doc, number, question, answer):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    run = heading.add_run(f"{number}. {question}")
    run.bold = True

    answer_p = doc.add_paragraph()
    answer_p.style = doc.styles["Normal"]
    answer_p.add_run("Suggested answer: ").bold = True
    answer_p.add_run(answer)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Normal"].font.color.rgb = RGBColor(31, 41, 55)

styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(22)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(17, 24, 39)

styles["Subtitle"].font.name = "Arial"
styles["Subtitle"].font.size = Pt(12)
styles["Subtitle"].font.color.rgb = RGBColor(75, 85, 99)

styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(17, 24, 39)

styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(31, 41, 55)

header = section.header.paragraphs[0]
header.text = "Digital Divide Oral Defense Q&A"
header.style = styles["Normal"]
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = RGBColor(107, 114, 128)
add_bottom_border(header)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("Page ")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(107, 114, 128)
fld_char_1 = OxmlElement("w:fldChar")
fld_char_1.set(qn("w:fldCharType"), "begin")
instr_text = OxmlElement("w:instrText")
instr_text.text = "PAGE"
fld_char_2 = OxmlElement("w:fldChar")
fld_char_2.set(qn("w:fldCharType"), "end")
footer_run._r.append(fld_char_1)
footer_run._r.append(instr_text)
footer_run._r.append(fld_char_2)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Oral Defense Q&A")

subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Bridging the Digital Divide Must Be a National Priority")

authors = doc.add_paragraph(style="Subtitle")
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run("Prepared for Jay Christian Estampa, Mariz Mae Pareja, and Josephine Vidal")

doc.add_paragraph()

meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = "Table Grid"
meta_table.autofit = False
for row in meta_table.rows:
    set_cell_width(row.cells[0], 2300)
    set_cell_width(row.cells[1], 7060)
for cell in meta_table.rows[0].cells:
    set_cell_shading(cell, "EEF2F7")
meta_table.cell(0, 0).text = "Editorial Type"
meta_table.cell(0, 1).text = "Persuasive editorial"
meta_table.cell(1, 0).text = "Main Stand"
meta_table.cell(1, 1).text = (
    "Bridging the digital divide must become a national priority because internet access now affects education, employment, governance, communication, and equal opportunity."
)
for row in meta_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading("Short Opening Statement", level=1)
doc.add_paragraph(
    "Good day. Our editorial argues that bridging the digital divide must be treated as a national priority. "
    "Internet access is no longer a luxury because education, employment, governance, and communication now depend "
    "on digital platforms. When rural and underserved communities lack reliable connectivity, they are also denied "
    "equal opportunities. Our stand is that the government, schools, private sector, and telecommunication companies "
    "must work together to improve infrastructure, reduce costs, provide public access points, and strengthen digital literacy."
)

doc.add_heading("Possible Oral Defense Questions", level=1)

qa = [
    (
        "What is the main argument of your editorial?",
        "Our main argument is that the digital divide must be addressed as a national priority because internet access now affects education, work, communication, public services, and equal opportunity.",
    ),
    (
        "Why did you choose this topic?",
        "We chose this topic because many people still experience limited or unreliable connectivity, especially in rural and underserved areas. This issue affects students, workers, families, and communities in everyday life.",
    ),
    (
        "Why is your editorial persuasive?",
        "It is persuasive because it does not only explain the problem; it takes a clear stand and urges action. We argue that the government and other sectors must actively bridge the digital divide through infrastructure, policy, and education.",
    ),
    (
        "What is the digital divide?",
        "The digital divide is the gap between people who have access to reliable internet and digital tools and those who do not. It also includes differences in digital skills and the ability to use technology effectively.",
    ),
    (
        "Why is internet access no longer a luxury?",
        "Internet access is no longer a luxury because many essential activities now depend on it, including online learning, job applications, business, communication, government services, research, and access to information.",
    ),
    (
        "Who is most affected by the digital divide?",
        "Rural communities, low-income families, students, and marginalized groups are often the most affected because they may lack stable internet, devices, affordable service, or digital literacy support.",
    ),
    (
        "How does the digital divide affect education?",
        "It affects education by making it harder for students to attend online classes, access learning materials, submit requirements, communicate with teachers, and develop digital skills needed for the future.",
    ),
    (
        "How does the digital divide affect the economy?",
        "It limits economic opportunities because people without internet access may struggle to find jobs, join online work, run digital businesses, access training, or participate in modern economic activities.",
    ),
    (
        "Why should the government lead the solution?",
        "The government should lead because the problem is large-scale and affects national development. Policies, funding, infrastructure programs, and public services are needed to ensure that access reaches underserved communities.",
    ),
    (
        "What specific solutions does your editorial propose?",
        "Our editorial proposes improving digital infrastructure, expanding network coverage, making internet service more affordable, providing free Wi-Fi in schools and public spaces, developing digital literacy, and encouraging cooperation among government, schools, private companies, and telecommunications providers.",
    ),
    (
        "Why are schools important in bridging the digital divide?",
        "Schools are important because they can provide access to digital tools and teach students how to use technology responsibly and effectively. Digital literacy helps students become more prepared for education, work, and civic participation.",
    ),
    (
        "Why is collaboration with the private sector necessary?",
        "Collaboration is necessary because telecommunication companies and private organizations have resources, technology, and technical expertise that can help expand connectivity and improve service delivery.",
    ),
    (
        "What is the strongest evidence you used?",
        "One strong point is that recent studies describe digital infrastructure as essential to education and institutional function. This supports our claim that internet access is now part of equal participation in society.",
    ),
    (
        "What is a possible weakness of your editorial?",
        "A possible weakness is that the editorial presents broad solutions that may require large budgets and long-term implementation. However, the urgency of the issue justifies strong national action and cooperation across sectors.",
    ),
    (
        "How would you respond to people who say internet access is not the government's responsibility?",
        "We would say that internet access is now connected to education, public services, employment, and national development. Because it affects equal opportunity, the government has a responsibility to help make access fairer and more inclusive.",
    ),
    (
        "Why is this issue urgent?",
        "It is urgent because every year without reliable access widens the gap between connected and disconnected communities. Delayed action can deepen inequality in education, employment, and social participation.",
    ),
    (
        "What is your conclusion in simple terms?",
        "Our conclusion is that bridging the digital divide is not only about technology. It is about fairness, opportunity, and inclusive growth. A nation cannot fully progress if many of its people remain disconnected.",
    ),
]

for index, (question, answer) in enumerate(qa, start=1):
    add_question_block(doc, index, question, answer)

doc.add_heading("Answering Strategy", level=1)
strategy = doc.add_paragraph()
strategy.add_run("If the panel asks why this should be a national priority, answer with this pattern: ").bold = True
strategy.add_run(
    "\"Because digital access now affects basic opportunities. When people are disconnected, they are also limited in education, work, services, and participation in society.\""
)

doc.add_heading("Quick Closing Statement", level=1)
doc.add_paragraph(
    "In conclusion, our editorial calls for urgent and collective action because the digital divide affects more than internet connection. It affects equality, education, livelihood, and national progress. Bridging this gap means giving more people a fair chance to participate in modern society."
)

doc.save(OUTPUT)
print(OUTPUT)
