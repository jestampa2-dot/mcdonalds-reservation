from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
OUT = ROOT / "Defense Q and A - McDonalds Event Planner.docx"
DARK = RGBColor(31, 31, 31)


QA = [
    (
        "1. What is the main problem your system solves?",
        "Our system solves the problem of manual event reservation handling, such as scheduling conflicts, delayed confirmation, incomplete customer records, and difficulty monitoring booking status. It provides one digital platform where customers, staff, and administrators can manage reservations more efficiently.",
    ),
    (
        "2. Why is your system better than manual reservation?",
        "It is better because it centralizes all reservation information in one system. Customers can book online, upload payment proof, and track their booking status, while administrators can review reservations, assign staff, and monitor schedules faster. This reduces errors, saves time, and improves communication.",
    ),
    (
        "3. What makes your system different from existing systems?",
        "Unlike general hotel, restaurant, or facility reservation systems, our system is designed specifically for McDonald's event reservations. It includes event packages, branch availability, menu selection, payment proof upload, admin approval, staff check-in, service status updates, and an Android mobile app for easier access.",
    ),
    (
        "4. Why did you choose Laravel, Vue, and React Native?",
        "We chose Laravel because it provides secure backend features such as routing, authentication, validation, database management, and file upload handling. We used Vue.js because it helps create a responsive and interactive web interface for customers, staff, and administrators. We chose React Native with Expo because it allows us to develop an Android mobile application that connects to the same Laravel backend.",
    ),
    (
        "5. How do you prevent double booking?",
        "The system prevents double booking by checking the selected branch, date, time, duration, and availability before accepting a reservation. If a schedule is already full or unavailable, the system can prevent users from choosing that slot or help the admin identify possible conflicts.",
    ),
    (
        "6. Explain your database structure.",
        "The database is organized around the Reservation table, which stores booking details such as customer, branch, event date, time, guests, package, payment proof, total amount, booking status, and service status. It is connected to the Users table for customer, staff, and admin accounts, the Branches table for branch details, and catalog-related tables such as Event Types, Booking Packages, Room Options, Menu Items, and Add-ons. This structure keeps all booking, user, branch, and package information centralized.",
    ),
    (
        "7. What is the gap between your project and your RRL?",
        "The gap is that most related studies focus on general reservation systems, hotel booking, school facility reservation, restaurant table booking, or catering services. They improve scheduling and record management, but they are not specifically designed for fast-food event reservations. Our project fills that gap by creating a system specifically for McDonald's event services, with features such as event packages, branch availability, menu selection, payment proof upload, admin approval, staff check-in, service status updates, and an Android mobile app.",
    ),
    (
        "8. What is your project all about?",
        "Our project is the McDonald's Event Planner, a digital reservation system for birthday parties, business meetings, and other reservations in McDonald's Surigao Branch. It allows customers to book events, upload payment proof, and track reservation status. It also helps staff and administrators manage bookings, check-ins, schedules, branches, packages, and reports.",
    ),
    (
        "9. Why did you choose this topic?",
        "We chose this topic because event reservations require organized scheduling, accurate records, and fast communication between customers and staff. A digital system can make the process more convenient for customers and easier to manage for the branch.",
    ),
    (
        "10. Who are the users of the system?",
        "The main users are customers, staff, and administrators. Customers create and monitor reservations. Staff handle check-ins and service updates. Administrators approve bookings, assign crew, manage branches, packages, accounts, and reports.",
    ),
    (
        "11. What methodology did you use?",
        "We used the Rapid Application Development or RAD methodology. RAD was suitable because our project required quick prototyping, user feedback, and continuous improvement of the system design and features.",
    ),
    (
        "12. Why RAD and not Waterfall?",
        "RAD is more flexible than Waterfall. Since our system involves user interfaces, booking flows, and feedback from possible users, RAD allowed us to revise the design and improve features faster instead of waiting until the end of development.",
    ),
    (
        "13. What are the main features of your system?",
        "The main features include customer registration and login, event booking, branch and schedule selection, package and menu selection, payment proof upload, booking dashboard, reschedule and cancellation, admin booking approval, staff check-in, service status update, reports, and account management.",
    ),
    (
        "14. Why did you create an Android mobile app?",
        "We created an Android mobile app because many users prefer using smartphones for online transactions. The app makes it easier for customers to book events, view dashboards, manage accounts, and access reservation details conveniently.",
    ),
    (
        "15. How does your booking process work?",
        "The customer logs in, selects an event type, branch, date, time, package, menu items, add-ons, and uploads payment proof. The system stores the booking as pending. The admin reviews it and updates the status. Staff can then prepare for the event and check in the customer.",
    ),
    (
        "16. How is payment handled?",
        "The current system allows customers to upload proof of payment. The administrator reviews the uploaded proof before confirming or updating the reservation status.",
    ),
    (
        "17. Does your system have online payment integration?",
        "Not yet. Online payment integration is part of our recommendation for future improvement. The current version uses payment proof upload.",
    ),
    (
        "18. How do you secure the system?",
        "The system uses authentication, role-based access, protected routes, and Laravel Sanctum for mobile API tokens. This ensures that customers, staff, and administrators can only access features allowed for their roles.",
    ),
    (
        "19. What is the role of the administrator?",
        "The administrator reviews pending bookings, approves or rejects reservations, assigns staff, manages branches, updates packages and catalog items, monitors availability, manages accounts, and views reports.",
    ),
    (
        "20. What is the role of staff?",
        "Staff can view today's bookings, check in guests using booking codes, update service status, and record service adjustments during event operations.",
    ),
    (
        "21. What is the role of customers?",
        "Customers can register, log in, create event reservations, upload payment proof, view booking status, reschedule, cancel, and manage their profile.",
    ),
    (
        "22. What are the expected benefits of your system?",
        "The system improves booking convenience, reduces manual errors, minimizes scheduling conflicts, centralizes records, speeds up confirmation, and helps staff and administrators monitor reservations more effectively.",
    ),
    (
        "23. What evaluation method did you use?",
        "The system was evaluated using selected ISO 25010 software quality criteria such as functional suitability, usability, performance efficiency, reliability, and security.",
    ),
    (
        "24. What are the limitations of your system?",
        "The system currently focuses on McDonald's event reservation management. It does not yet include iOS support, direct online payment gateway integration, advanced analytics, or full offline functionality.",
    ),
    (
        "25. What are your recommendations?",
        "We recommend adding iOS compatibility, online payment integration, push notifications, SMS or email updates, improved QR check-in, advanced reports, audit logs, and wider testing with more users.",
    ),
    (
        "26. Did your project meet its objectives?",
        "Yes. The project met its objectives by developing a web and Android reservation system that supports customer booking, staff operations, administrator management, centralized records, and mobile access.",
    ),
    (
        "27. What is the most important contribution of your project?",
        "The most important contribution is providing an integrated reservation platform specifically designed for McDonald's event services. It connects customers, staff, and administrators in one organized system.",
    ),
    (
        "28. What will happen if the internet connection is unstable?",
        "Some mobile screens use cached data, so recently loaded information may still be visible. However, actions such as submitting reservations, uploading payment proof, and updating booking status still require internet connection.",
    ),
    (
        "29. Can the system be used by other branches?",
        "Yes, the system can be expanded for other branches. It already includes branch management concepts, but future improvement may include stronger multi-branch deployment and branch-specific settings.",
    ),
    (
        "30. If you were given more time, what would you improve first?",
        "We would prioritize online payment integration and real-time notifications because these would make the reservation process faster and more convenient for customers and administrators.",
    ),
]


def set_run(run, size=12, bold=False, color=DARK):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Defense Q&A")
    set_run(run, size=18, bold=True, color=RGBColor(191, 31, 36))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("McDonald's Event Planner Capstone Project")
    set_run(run, size=12, bold=True)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run("Use this document to practice concise, defense-ready answers. You can memorize the main idea of each answer instead of memorizing every word.")
    set_run(run, size=11)

    for question, answer in QA:
        q = doc.add_paragraph()
        q.paragraph_format.space_before = Pt(8)
        q.paragraph_format.space_after = Pt(3)
        run = q.add_run(question)
        set_run(run, size=12, bold=True, color=RGBColor(191, 31, 36))

        a = doc.add_paragraph()
        a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        a.paragraph_format.space_after = Pt(4)
        run = a.add_run(answer)
        set_run(run, size=11)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_doc()
