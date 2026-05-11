from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\User\OneDrive\Documents\Cpastone Project\Chapter2.docx")
OUTPUT = Path(r"C:\Users\User\mcdonalds-reservation\Chapter2 - Organized Numbered Sections.docx")

HEADINGS_BEFORE_PARAGRAPH = {
    1: "2.1 Overview of Related Literature",
    2: "2.2 Summary of Local Reservation and Hospitality Studies",
    5: "2.3 Summary of Foreign Digital Reservation Studies",
    9: "2.5 Local Literature",
    11: "2.5.1 Centralized Booking and Availability Management",
    14: "2.5.2 Institutional, Hospitality, and Management Information Systems",
    16: "2.5.3 User Experience, Web-Based Access, and Mobile Compatibility",
    27: "2.6.1 Online Event Booking and Hospitality Platforms",
    29: "2.6.2 Digital Transformation and ICT Integration",
    32: "2.6.3 Web-Based Restaurant and Catering Reservation Systems",
    36: "2.6.4 Mobile Access, Automation, Capacity, and Emerging Technologies",
}

REPLACE_PARAGRAPH_TEXT = {
    6: "2.4 Gap in the Literature",
    25: "2.6 Foreign Literature",
}


def format_heading(paragraph, level=2):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(10 if level == 2 else 8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(31, 31, 31)


def insert_heading_before(paragraph, text):
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    new_paragraph.add_run(text)
    level = 3 if text.count(".") >= 2 else 2
    format_heading(new_paragraph, level)


def replace_text(paragraph, text):
    paragraph.text = ""
    paragraph.add_run(text)
    level = 3 if text.count(".") >= 2 else 2
    format_heading(paragraph, level)


def organize():
    doc = Document(str(SOURCE))
    original_paragraphs = list(doc.paragraphs)

    for index, text in REPLACE_PARAGRAPH_TEXT.items():
        if index < len(original_paragraphs):
            replace_text(original_paragraphs[index], text)

    for index in sorted(HEADINGS_BEFORE_PARAGRAPH.keys(), reverse=True):
        if index < len(original_paragraphs):
            insert_heading_before(original_paragraphs[index], HEADINGS_BEFORE_PARAGRAPH[index])

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    organize()
