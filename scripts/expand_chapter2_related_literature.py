from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\User\OneDrive\Documents\Cpastone Project\Chapter2.docx")
OUTPUT = Path(r"C:\Users\User\mcdonalds-reservation\Chapter2 - Expanded Related Literature.docx")


EXPANSIONS = {
    1: [
        "For the proposed McDonald's Event Planner, these studies are important because they show how digital systems can improve the flow of reservation-related activities from inquiry to confirmation. The literature also supports the idea that reservation systems should not only record bookings, but also help users check availability, manage customer details, monitor status updates, and support staff decision-making. In this study, the related literature serves as a basis for designing a system that connects customer booking, administrator approval, and branch-level event preparation.",
    ],
    2: [
        "These local findings are relevant to the proposed system because restaurant and event reservations require accurate scheduling and organized records. When reservations are handled manually, staff may encounter problems such as duplicate bookings, incomplete customer information, unclear payment status, and delayed confirmation. Automation can reduce these problems by placing reservation information in a single system where authorized users can review, update, and monitor booking records.",
    ],
    3: [
        "Centralized booking is especially useful for a business with several reservation details to manage, such as event type, branch, date, time, number of guests, room choice, package, menu items, add-ons, and payment proof. A centralized system gives customers and administrators a shared source of information, which helps avoid confusion between what the customer submitted and what the staff must prepare. This concept directly supports the proposed McDonald's Event Planner because it needs a unified reservation record for customer, staff, and administrator users.",
    ],
    4: [
        "Although many local reservation systems are designed for schools and facilities, their concepts can still be applied to food-service event booking. Availability checking, user authentication, database storage, and administrative approval are common requirements across different reservation systems. The difference is that the proposed study applies these concepts to McDonald's event reservations, where food packages, branch capacity, payment proof, and staff preparation are part of the booking workflow.",
    ],
    5: [
        "The foreign studies also show that reservation systems are becoming more user-centered. Customers expect faster confirmation, clearer status updates, and convenient access through web or mobile platforms. These expectations are reflected in the proposed system through features such as an online booking form, customer dashboard, mobile access, and administrator review tools.",
    ],
    7: [
        "This gap is significant because event reservations in a quick-service restaurant setting involve both customer-facing and operations-facing requirements. Customers need a simple way to book and track events, while staff and administrators need accurate information for approval, scheduling, preparation, and service updates. The proposed system addresses this gap by combining booking, catalog selection, payment proof upload, admin review, staff check-in, and mobile access in one reservation workflow.",
    ],
    9: [
        "Their findings indicate that online reservation tools can improve service delivery by making information available to both customers and management. For the proposed system, this supports the inclusion of an online booking process where customer details and reservation preferences are stored immediately after submission. This reduces the dependence on handwritten forms or repeated manual communication.",
    ],
    10: [
        "This finding is applicable to McDonald's event reservations because small and medium service operations often need practical systems that can organize customer requests without requiring complex enterprise software. A computerized system can help staff view bookings more clearly, check pending requests, and update the status of reservations. As a result, the customer receives a more reliable booking experience.",
    ],
    11: [
        "Real-time availability tracking is important in the proposed system because an event reservation depends on the branch schedule and capacity. If customers and administrators can view available dates and time slots, the chance of overlapping reservations is reduced. This also helps administrators make faster decisions when confirming or rescheduling bookings.",
    ],
    14: [
        "The same principle applies to restaurant event reservations because staff members need a convenient way to identify which events are scheduled for a specific date and branch. A digital reservation system can reduce the workload of checking multiple records manually. It can also provide a clearer basis for planning room use, food preparation, and staff assignment.",
    ],
    15: [
        "The use of MIS concepts is relevant because the proposed system is not limited to accepting bookings. It also supports management decisions through reports, status monitoring, branch information, and reservation records. This means that the system can help administrators understand booking activity and prepare operational responses based on stored data.",
    ],
    16: [
        "These factors are important in designing the mobile and web interfaces of the proposed system. If the booking process is too complicated, users may abandon the reservation or provide incomplete information. Therefore, the system uses guided forms, clear labels, status indicators, and accessible customer dashboard features to improve user experience.",
    ],
    17: [
        "The limitation identified in their study supports the need for a system that is more specific to event booking. In the proposed McDonald's Event Planner, reservation details are connected with event packages, branch schedules, staff preparation, and payment proof. This makes the system more aligned with the actual requirements of restaurant-based event planning.",
    ],
    18: [
        "Data accessibility is one of the most important benefits of a web-based reservation system. When reservation details are available through a shared system, customers can monitor their bookings while administrators and staff can update operational information. This improves coordination among all users involved in the event process.",
    ],
    19: [
        "Automated scheduling also improves the accuracy of reservation decisions. In the proposed system, schedule-related information such as date, time, duration, and branch capacity must be checked before a booking can be treated as valid. This reduces the possibility of accepting reservations that cannot be accommodated by the selected branch.",
    ],
    20: [
        "Customer experience is improved when users receive clear confirmation and status information. This is why the proposed system includes booking status labels such as pending, confirmed, cancelled, rescheduled, and completed. These status updates help customers understand the progress of their reservation and reduce the need for repeated follow-up messages.",
    ],
    21: [
        "Mobile compatibility is particularly relevant to the proposed study because many customers prefer to complete transactions through smartphones. A mobile-friendly reservation process allows users to book events, upload payment proof, and check dashboard updates without needing to use a desktop computer. This supports the inclusion of the Android mobile application as part of the system output.",
    ],
    26: [
        "These foreign studies provide a broader view of how digital reservation systems are used in service industries outside the local context. They show that online booking has become an expected feature in hospitality and food-service operations because it improves convenience and helps businesses organize demand. The proposed system adopts these ideas by providing both web and mobile access to reservation functions.",
    ],
    27: [
        "Their study is relevant because it emphasizes remote access, which is also a major requirement of the proposed system. Customers should be able to submit reservation requests without visiting the branch first. This convenience can increase customer engagement while giving administrators a structured record to review.",
    ],
    28: [
        "Ease of use and real-time booking are important design considerations for the proposed mobile application. The system must present reservation choices in a way that customers can understand quickly, such as event type, branch, schedule, package, and payment proof. A clear interface can improve satisfaction and reduce errors during booking submission.",
    ],
    29: [
        "This supports the integration of reservation data with customer and operational records in the proposed system. Instead of treating a booking as an isolated transaction, the system connects it to the customer account, branch, package, payment proof, and staff service status. This creates a more complete view of the reservation lifecycle.",
    ],
    30: [
        "ICT-based systems also allow businesses to respond faster to customer needs. In the proposed system, administrators can review pending bookings, assign staff, update reservation status, and monitor availability through the dashboard. This improves coordination between customer requests and branch operations.",
    ],
    32: [
        "This recent literature supports the use of web-based tools for organizing food-service reservations. However, the proposed McDonald's Event Planner extends the idea by combining event booking with menu selection, branch availability, payment proof upload, and staff preparation features. These additions make the system more suitable for restaurant event management.",
    ],
    33: [
        "The catering context is related to the proposed system because both involve food preparation, scheduling, and customer-specific event details. The difference is that the proposed study focuses on McDonald's branch events, where the booking must consider branch capacity, event packages, room options, and operational staff actions.",
    ],
    34: [
        "This study supports the need to automate table and event booking in restaurant settings. Manual reservation processes may lead to lost records or delayed responses, especially when several customers inquire at the same time. A computerized system can improve accuracy by storing each reservation in a database and making it available for review.",
    ],
    35: [
        "The use of full-stack web technologies in restaurant booking demonstrates that modern reservation systems can be built using modular frontend and backend tools. In the proposed system, Laravel, Vue.js, and Expo React Native serve a similar purpose by separating the interface, business logic, and mobile access while keeping the data centralized.",
    ],
    36: [
        "Mobile-based reservation systems are strongly related to the proposed Android application. The mobile app gives customers easier access to booking features and gives staff or administrators selected operational tools. This supports the idea that mobile access can improve both convenience and engagement.",
    ],
    37: [
        "Automation can also help improve customer flow by reducing waiting time and making reservations more predictable. For event reservations, this means that staff can prepare food, room setup, and service requirements ahead of time. The proposed system supports this by providing staff with booking and service information before the event begins.",
    ],
    38: [
        "Pre-ordering is related to the proposed system's menu and add-on selection features. Allowing customers to choose food bundles and manual menu items during the booking process helps the branch prepare more accurately. This also gives administrators a clearer view of the expected event requirements.",
    ],
    39: [
        "Improved accessibility is also reflected in the proposed system's mobile-first design. Customers can access the booking flow through a phone, while authenticated users can view dashboard information and account details. This makes the reservation process more convenient for users who rely on mobile devices.",
    ],
    40: [
        "Capacity optimization is important for restaurant event reservations because each branch has limits on guests, time slots, and concurrent events. The proposed system applies this concept through branch capacity, schedule availability, and booking settings. These controls help administrators avoid overbooking and improve resource planning.",
    ],
    41: [
        "Although the proposed system does not focus on artificial intelligence, the study supports the general direction of using automation to improve service operations. Automated booking records, status updates, and staff dashboards can already provide significant efficiency improvements. Future versions of the system may include predictive reports or recommendation features.",
    ],
    42: [
        "Innovation in restaurant business systems shows that digital platforms are now part of competitive service delivery. The proposed system contributes to this direction by modernizing event reservations through web and mobile access. This helps align restaurant operations with current customer expectations for fast and accessible digital services.",
    ],
    43: [
        "Reception automation is relevant because event reservations often begin with customer inquiry and confirmation. By automating parts of this process, the proposed system can reduce repetitive staff tasks and provide customers with a clearer booking path. This supports the study's goal of making event reservation management more organized and efficient.",
    ],
}


def style_like(paragraph, reference):
    paragraph.style = reference.style
    paragraph.alignment = reference.alignment
    paragraph.paragraph_format.first_line_indent = reference.paragraph_format.first_line_indent
    paragraph.paragraph_format.space_before = reference.paragraph_format.space_before
    paragraph.paragraph_format.space_after = reference.paragraph_format.space_after
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)


def insert_paragraph_after(paragraph, text):
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    run = new_paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    style_like(new_paragraph, paragraph)
    return new_paragraph


def expand_document():
    doc = Document(str(SOURCE))
    original_paragraphs = list(doc.paragraphs)

    for index in sorted(EXPANSIONS.keys(), reverse=True):
        if index >= len(original_paragraphs):
            continue
        anchor = original_paragraphs[index]
        for text in reversed(EXPANSIONS[index]):
            insert_paragraph_after(anchor, text)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    expand_document()
