from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
TEMPLATE = Path(r"C:\Users\User\Downloads\BSIT Capstone Project Template.docx")
OUT = ROOT / "Chapter 5 - Results and Discussion.docx"
USE_CASE = ROOT / "scripts" / "chapter5_use_case.png"
CLASS_DIAGRAM = ROOT / "scripts" / "chapter5_class_diagram.png"
SEQUENCE_DIAGRAM = ROOT / "scripts" / "chapter5_sequence_diagram.png"
MOBILE_UI = ROOT / "scripts" / "chapter5_mobile_app_design.png"


ACCENT = RGBColor(191, 31, 36)
DARK = RGBColor(31, 31, 31)
HEADER = "F4E1C1"
LIGHT = "FFF7E6"
BULLET_NUM_ID = None


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.2)
    run.bold = bold
    run.font.color.rgb = DARK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(tag_name)):
        value = element.get(qn(attr_name))
        if value and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def ensure_bullet_numbering(doc) -> int:
    global BULLET_NUM_ID
    if BULLET_NUM_ID is not None:
        return BULLET_NUM_ID

    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "\u2022"), ("w:lvlJc", "left")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        lvl.append(el)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    BULLET_NUM_ID = num_id
    return num_id


def add_para(doc, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return paragraph


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = ACCENT if level == 1 else DARK
    return paragraph


def add_bullet(doc, text: str):
    num_id = ensure_bullet_numbering(doc)
    paragraph = doc.add_paragraph(style="Normal")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK
    return paragraph


def add_caption(doc, text: str):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(9)
    run = caption.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.italic = True


def add_image(doc, path: Path, width: float, title: str, alt: str, caption: str):
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline_shape = paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", alt)
    add_caption(doc, caption)


def draw_rounded_box(draw, coords, title, body, font_title, font_body, fill=(255, 247, 230)):
    draw.rounded_rectangle(coords, radius=18, fill=fill, outline=(184, 184, 184), width=3)
    x1, y1, x2, _ = coords
    draw.multiline_text((x1 + 22, y1 + 18), title, fill=(31, 31, 31), font=font_title, spacing=4)
    if body:
        draw.multiline_text((x1 + 22, y1 + 64), body, fill=(55, 55, 55), font=font_body, spacing=5)


def arrow(draw, start, end, color=(191, 31, 36), width=5):
    draw.line((start, end), fill=color, width=width)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        points = [(ex, ey), (ex - 18 if ex >= sx else ex + 18, ey - 11), (ex - 18 if ex >= sx else ex + 18, ey + 11)]
    else:
        points = [(ex, ey), (ex - 11, ey - 18 if ey >= sy else ey + 18), (ex + 11, ey - 18 if ey >= sy else ey + 18)]
    draw.polygon(points, fill=color)


def load_fonts():
    try:
        return (
            ImageFont.truetype("arialbd.ttf", 36),
            ImageFont.truetype("arialbd.ttf", 24),
            ImageFont.truetype("arial.ttf", 20),
            ImageFont.truetype("arial.ttf", 18),
        )
    except Exception:
        font = ImageFont.load_default()
        return font, font, font, font


def create_use_case(path: Path):
    if Image is None:
        return
    title_font, box_font, small_font, caption_font = load_fonts()
    img = Image.new("RGB", (1500, 900), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((28, 28, 1472, 872), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "Use-Case Diagram of the McDonald's Reservation System", fill=(31, 31, 31), font=title_font)
    draw.rounded_rectangle((315, 145, 1185, 760), radius=28, outline=(191, 31, 36), width=4, fill=(255, 253, 248))
    draw.text((585, 165), "Reservation System", fill=(31, 31, 31), font=box_font)

    actors = [("Customer", 105, 305), ("Staff", 105, 575), ("Admin", 1280, 440)]
    for label, x, y in actors:
        draw.ellipse((x, y, x + 50, y + 50), outline=(31, 31, 31), width=3)
        draw.line((x + 25, y + 50, x + 25, y + 130), fill=(31, 31, 31), width=3)
        draw.line((x - 20, y + 80, x + 70, y + 80), fill=(31, 31, 31), width=3)
        draw.line((x + 25, y + 130, x - 15, y + 190), fill=(31, 31, 31), width=3)
        draw.line((x + 25, y + 130, x + 65, y + 190), fill=(31, 31, 31), width=3)
        draw.text((x - 20, y + 205), label, fill=(31, 31, 31), font=box_font)

    cases = [
        ("Register / Login", 390, 240), ("Create Reservation", 610, 240), ("Upload Payment Proof", 850, 240),
        ("Track / Reschedule\nBooking", 500, 390), ("Generate QR Pass", 760, 390),
        ("Check In Guest", 440, 575), ("Update Service\nStatus", 700, 575),
        ("Review Bookings", 950, 395), ("Manage Catalog,\nBranches, Accounts", 905, 570),
    ]
    centers = {}
    for text, x, y in cases:
        draw.ellipse((x, y, x + 190, y + 85), fill=(255, 247, 230), outline=(184, 184, 184), width=3)
        draw.multiline_text((x + 26, y + 24), text, fill=(31, 31, 31), font=small_font, spacing=3)
        centers[text] = (x + 95, y + 42)

    for target in [cases[0], cases[1], cases[2], cases[3], cases[4]]:
        draw.line((180, 410, target[1] + 95, target[2] + 42), fill=(120, 120, 120), width=2)
    for target in [cases[5], cases[6]]:
        draw.line((180, 680, target[1] + 95, target[2] + 42), fill=(120, 120, 120), width=2)
    for target in [cases[7], cases[8], cases[6]]:
        draw.line((1280, 545, target[1] + 95, target[2] + 42), fill=(120, 120, 120), width=2)

    draw.text((60, 805), "Figure 5.1. Use-case diagram showing the interactions of customer, staff, and administrator users.", fill=(80, 80, 80), font=caption_font)
    img.save(path, quality=95)


def create_class_diagram(path: Path):
    if Image is None:
        return
    title_font, box_font, small_font, caption_font = load_fonts()
    img = Image.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((28, 28, 1472, 952), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "Class Diagram of the Reservation System", fill=(31, 31, 31), font=title_font)

    boxes = [
        (90, 150, 390, 370, "User", "id\nname\nemail\nphone\nrole\npassword"),
        (585, 130, 915, 425, "Reservation", "id\nuser_id\nbranch_code\nevent_date\nevent_time\nduration_hours\nguests\ntotal_amount\nstatus\nservice_status"),
        (1110, 150, 1410, 390, "Branch", "id\ncode\nname\ncity\nconcurrent_limit\nmax_guests\nis_active"),
        (90, 545, 390, 800, "BookingPackage", "id\nevent_type_id\ncode\nname\nprice\nguest_range\nis_active"),
        (585, 560, 915, 790, "MenuItem / AddOn", "id\ncategory_id\ncode\nname\nprice\nis_active"),
        (1110, 545, 1410, 770, "BookingSetting", "id\nopening_hour\nclosing_hour\ndefault_duration\ncapacity rules"),
    ]
    for box in boxes:
        x1, y1, x2, y2, title, body = box
        draw_rounded_box(draw, (x1, y1, x2, y2), title, body, box_font, caption_font)
        draw.line((x1, y1 + 55, x2, y1 + 55), fill=(184, 184, 184), width=2)

    arrow(draw, (390, 250), (585, 250), color=(90, 90, 90), width=3)
    arrow(draw, (1110, 250), (915, 250), color=(90, 90, 90), width=3)
    arrow(draw, (240, 545), (585, 410), color=(90, 90, 90), width=3)
    arrow(draw, (750, 560), (750, 425), color=(90, 90, 90), width=3)
    arrow(draw, (1260, 545), (1260, 390), color=(90, 90, 90), width=3)

    draw.text((405, 225), "creates", fill=(70, 70, 70), font=caption_font)
    draw.text((955, 225), "assigned to", fill=(70, 70, 70), font=caption_font)
    draw.text((285, 455), "selected package", fill=(70, 70, 70), font=caption_font)
    draw.text((780, 465), "menu selections", fill=(70, 70, 70), font=caption_font)
    draw.text((1100, 440), "uses rules", fill=(70, 70, 70), font=caption_font)
    draw.text((60, 890), "Figure 5.2. Class diagram showing the main data entities used by the reservation system.", fill=(80, 80, 80), font=caption_font)
    img.save(path, quality=95)


def create_sequence_diagram(path: Path):
    if Image is None:
        return
    title_font, box_font, small_font, caption_font = load_fonts()
    img = Image.new("RGB", (1500, 980), "white")
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((28, 28, 1472, 952), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "Sequence Diagram for Creating a Reservation", fill=(31, 31, 31), font=title_font)

    participants = [("Customer", 150), ("Mobile/Web UI", 430), ("Laravel API", 710), ("Database", 990), ("Admin", 1270)]
    for label, x in participants:
        draw.rounded_rectangle((x - 95, 145, x + 95, 205), radius=14, fill=(255, 247, 230), outline=(184, 184, 184), width=3)
        draw.text((x - 70, 162), label, fill=(31, 31, 31), font=small_font)
        draw.line((x, 205, x, 820), fill=(180, 180, 180), width=3)

    steps = [
        (150, 430, 260, "1. Enter booking details"),
        (430, 710, 340, "2. Submit reservation request"),
        (710, 990, 420, "3. Validate and save records"),
        (990, 710, 500, "4. Return booking reference"),
        (710, 430, 580, "5. Show pending status"),
        (1270, 710, 660, "6. Review and update status"),
        (710, 990, 735, "7. Store final status"),
    ]
    for x1, x2, y, label in steps:
        arrow(draw, (x1, y), (x2, y), width=4)
        draw.text((min(x1, x2) + 20, y - 28), label, fill=(31, 31, 31), font=caption_font)

    draw.text((60, 890), "Figure 5.3. Sequence diagram showing how a reservation request moves from the user interface to the backend and database.", fill=(80, 80, 80), font=caption_font)
    img.save(path, quality=95)


def create_mobile_ui(path: Path):
    if Image is None:
        return
    title_font, box_font, small_font, caption_font = load_fonts()
    img = Image.new("RGB", (1600, 980), (255, 199, 44))
    draw = ImageDraw.Draw(img)
    draw.rounded_rectangle((28, 28, 1572, 952), radius=24, outline=(190, 152, 38), width=3, fill=(255, 199, 44))
    draw.text((60, 55), "Current Android Mobile Application Design", fill=(35, 22, 11), font=title_font)

    screens = [
        (90, 145, "Welcome!", "Choose a branch,\ntime, and package.", "HOME", ["Party meals", "Book Event\nFaster", "Featured packages"], "Home"),
        (440, 145, "Book Event", "Plan branch,\nschedule, and menu.", "BOOK", ["Event chips", "Calendar slots", "Upload payment\nproof"], "Book"),
        (790, 145, "My Dashboard", "Booking, payments,\nand event details.", "DASHBOARD", ["Upcoming", "Confirmed spend", "Pending approvals"], "Dash"),
        (1140, 145, "My Account", "Profile, security,\nand verification.", "ACCOUNT", ["Customer details", "Password", "Sign out"], "Acct"),
    ]

    for x, y, title, subtitle, label, items, active_tab in screens:
        draw.rounded_rectangle((x, y, x + 285, y + 650), radius=34, fill=(35, 22, 11), outline=(80, 70, 65), width=4)
        draw.rounded_rectangle((x + 12, y + 12, x + 273, y + 638), radius=28, fill=(255, 199, 44))

        draw.text((x + 30, y + 42), title, fill=(35, 22, 11), font=box_font)
        draw.multiline_text((x + 30, y + 76), subtitle, fill=(106, 86, 71), font=caption_font, spacing=3)
        draw.ellipse((x + 213, y + 38, x + 255, y + 80), fill=(218, 41, 28), outline=(158, 31, 18), width=2)
        draw.text((x + 225, y + 44), "M", fill=(255, 199, 44), font=box_font)

        draw.rounded_rectangle((x + 28, y + 130, x + 257, y + 535), radius=28, fill=(255, 248, 231), outline=(240, 217, 181), width=2)
        draw.text((x + 48, y + 158), label, fill=(139, 91, 0), font=caption_font)

        if active_tab == "Home":
            draw.rounded_rectangle((x + 48, y + 195, x + 237, y + 292), radius=22, fill=(255, 225, 142), outline=(240, 217, 181), width=2)
            draw.text((x + 68, y + 214), "Customer mobile", fill=(139, 91, 0), font=caption_font)
            draw.multiline_text((x + 68, y + 242), items[1], fill=(35, 22, 11), font=caption_font, spacing=2)
            card_y = y + 315
            for item in (items[0], items[2]):
                draw.rounded_rectangle((x + 48, card_y, x + 237, card_y + 58), radius=16, fill=(255, 253, 248), outline=(240, 217, 181), width=2)
                draw.text((x + 65, card_y + 18), item, fill=(35, 22, 11), font=caption_font)
                card_y += 74
        elif active_tab == "Book":
            cy = y + 198
            chip_colors = [(255, 225, 142), (248, 216, 216), (207, 242, 200)]
            for i, item in enumerate(items):
                draw.rounded_rectangle((x + 48, cy, x + 237, cy + 66), radius=18, fill=chip_colors[i], outline=(240, 217, 181), width=2)
                draw.multiline_text((x + 65, cy + 15), item, fill=(35, 22, 11), font=caption_font, spacing=2)
                cy += 82
            draw.rounded_rectangle((x + 48, y + 455, x + 237, y + 503), radius=20, fill=(218, 41, 28))
            draw.text((x + 83, y + 468), "Submit Booking", fill=(255, 255, 255), font=caption_font)
        elif active_tab == "Dash":
            cy = y + 195
            values = ["2", "P4,500", "1"]
            for item, value in zip(items, values):
                draw.rounded_rectangle((x + 48, cy, x + 237, cy + 68), radius=18, fill=(255, 253, 248), outline=(240, 217, 181), width=2)
                draw.text((x + 64, cy + 13), item, fill=(106, 86, 71), font=caption_font)
                draw.text((x + 64, cy + 36), value, fill=(35, 22, 11), font=small_font)
                cy += 82
            draw.rounded_rectangle((x + 48, y + 455, x + 237, y + 503), radius=18, fill=(255, 240, 190), outline=(240, 217, 181), width=2)
            draw.text((x + 68, y + 468), "Booking card", fill=(35, 22, 11), font=caption_font)
        else:
            cy = y + 195
            for item in items:
                draw.rounded_rectangle((x + 48, cy, x + 237, cy + 62), radius=18, fill=(255, 253, 248), outline=(240, 217, 181), width=2)
                draw.text((x + 65, cy + 20), item, fill=(35, 22, 11), font=caption_font)
                cy += 78
            draw.ellipse((x + 48, y + 440, x + 96, y + 488), fill=(255, 199, 44), outline=(35, 22, 11), width=2)
            draw.text((x + 63, y + 453), "U", fill=(35, 22, 11), font=small_font)

        draw.rounded_rectangle((x + 30, y + 555, x + 255, y + 608), radius=16, fill=(255, 247, 238), outline=(231, 222, 209), width=2)
        tabs = ["Home", "Book", "Dash", "Acct"]
        for i, tab in enumerate(tabs):
            tx = x + 44 + i * 53
            active = tab == active_tab
            if active:
                draw.rounded_rectangle((tx - 8, y + 565, tx + 43, y + 598), radius=10, fill=(246, 214, 214))
            draw.text((tx, y + 573), tab, fill=(35, 22, 11) if active else (125, 103, 85), font=caption_font)

    draw.text((60, 885), "Figure 5.4. Current mobile app design based on the Expo screens, yellow customer pages, McLogo headers, rounded cards, chips, and bottom tabs.", fill=(80, 65, 35), font=caption_font)
    img.save(path, quality=95)


def add_table(doc, title, headers, rows):
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], HEADER)
    mark_header_row(table.rows[0])
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT)


def build_doc():
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("5. RESULTS AND DISCUSSION")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = ACCENT

    add_para(doc, "This chapter presents the results of the development process for the proposed McDonald's Reservation System. The outputs include project planning results, system design diagrams, implemented modules, internal testing observations, and the Android mobile application design. The discussion explains how the system output responds to the reservation problems identified in the study and how the customer, staff, and administrator modules work together to support a more organized digital booking process.")
    add_para(doc, "The development output produced a web-based reservation management system and a mobile application connected through a Laravel backend. The system supports customer registration, login, booking creation, payment proof upload, reservation monitoring, rescheduling, cancellation, staff check-in, service status updates, administrator booking review, crew assignment, branch management, catalog management, availability monitoring, and account management. These results show that the proposed system can centralize reservation activities that are usually handled through scattered manual communication.")
    add_para(doc, "The results are discussed according to the major outputs of the Systems Analysis and Design process. These include project planning, system modeling, implementation, evaluation, and the Android mobile application. The diagrams and tables in this chapter summarize the completed system structure and show how data flows from the customer interface to the backend, database, staff operations, and administrator dashboard.")

    add_heading(doc, "5.1 Project Planning", 1)
    add_para(doc, "Project planning focused on organizing the development activities, identifying resources, and defining the expected outputs for each phase of the project. The Rapid Application Development approach guided the timeline because the researchers needed to analyze requirements, create prototypes, gather feedback, implement the software, and evaluate the system within a limited development period. The planning phase helped ensure that the work was divided into manageable tasks.")
    add_para(doc, "The project resources included software tools, hardware devices, development personnel, and testing participants. The software resources included Laravel, PHP, Vue.js, Inertia.js, Tailwind CSS, Expo, React Native, TypeScript, Figma, and a relational database. Hardware resources included a laptop or desktop computer for development and Android devices or emulators for mobile testing. Human resources included the researchers, possible customer users, staff users, and administrator representatives who could review the prototype and test the system.")
    add_table(
        doc,
        "Project Timeline and Gantt Summary",
        ["Phase", "Activities", "Expected Output"],
        [
            ("Week 1", "Problem identification, project scope definition, and requirements gathering.", "Approved topic, project objectives, and initial requirements list."),
            ("Week 2", "System analysis, user role definition, database planning, and process modeling.", "User roles, data entities, and system flow design."),
            ("Week 3", "Figma prototype preparation for customer, staff, and admin interfaces.", "Low-fidelity and mobile-first prototype screens."),
            ("Weeks 4-5", "Backend, web interface, mobile interface, API, and database implementation.", "Working Laravel web system and Expo mobile application."),
            ("Week 6", "Internal testing, corrections, documentation, and evaluation preparation.", "Tested system modules, evaluation tools, and final documentation."),
        ],
    )
    add_para(doc, "The timeline shows that the project followed an iterative development process. Some tasks overlapped because feedback from prototype review affected database fields, interface layout, and module priorities. This flexible planning helped the researchers revise the system design without waiting until the final stage of development.")
    add_table(
        doc,
        "Resource Allocation",
        ["Resource", "Description", "Use in the Project"],
        [
            ("Development Tools", "Laravel, Vue.js, Inertia.js, Tailwind CSS, Expo, React Native, TypeScript.", "Used to build the backend, web interface, and Android mobile app."),
            ("Design Tools", "Figma and design references from the mobile interface.", "Used to plan screen layout, navigation, and mobile app flow."),
            ("Database Resources", "Laravel migrations, Eloquent models, and reservation data tables.", "Used to store users, bookings, branches, catalog items, and status records."),
            ("Testing Resources", "Developer testing, selected users, mobile device or emulator, and browser testing.", "Used to verify functionality, usability, performance, and security behavior."),
            ("Documentation Resources", "Capstone template, system diagrams, and implementation notes.", "Used to prepare the technical paper and organize project outputs."),
        ],
    )

    add_heading(doc, "5.2 Systems Design", 1)
    add_para(doc, "The system design outputs show the structure and behavior of the proposed reservation system. The design models were prepared to clarify how users interact with the system, how the main data entities are related, and how a reservation request is processed. These diagrams support the SAD requirement that the proposed system must be understandable before and during implementation.")

    create_use_case(USE_CASE)
    add_heading(doc, "a.) Use-Case Diagram", 2)
    add_para(doc, "The use-case diagram illustrates the interactions of the three main actors: Customer, Staff, and Admin. Customers use the system to register, log in, create reservations, upload payment proof, track booking status, reschedule or cancel bookings, and access the QR or booking pass. Staff users handle event preparation, guest check-in, and service updates. Administrators review bookings, update reservation status, assign crew, manage catalog details, manage branches, monitor availability, and maintain user accounts.")
    add_image(doc, USE_CASE, 6.4, "Use-Case Diagram", "Use-case diagram showing customer, staff, and admin interactions with the reservation system.", "Figure 5.1. Use-Case Diagram of the McDonald's Reservation System")
    add_para(doc, "The use-case design shows that the system separates responsibilities according to user roles. This separation improves security and usability because customers only access booking-related functions, staff members access operational tools, and administrators access management functions.")

    create_class_diagram(CLASS_DIAGRAM)
    add_heading(doc, "b.) Class Diagram", 2)
    add_para(doc, "The class diagram presents the main entities used by the system. The central class is Reservation because it connects the customer, branch, package, menu selections, payment proof, booking status, and service status. The User class represents customer, staff, and administrator accounts. The Branch class stores branch-specific details such as code, location, capacity, and supported event types. Other classes such as BookingPackage, MenuItem, AddOn, and BookingSetting support the reservation options and rules used during booking.")
    add_image(doc, CLASS_DIAGRAM, 6.4, "Class Diagram", "Class diagram showing the User, Reservation, Branch, BookingPackage, MenuItem, AddOn, and BookingSetting entities.", "Figure 5.2. Class Diagram of the Reservation System")
    add_para(doc, "This design supports maintainability because each major data concept is represented by its own model. Changes to packages, branches, menu items, or booking settings can be handled without changing the entire reservation structure.")

    create_sequence_diagram(SEQUENCE_DIAGRAM)
    add_heading(doc, "c.) Sequence Diagram", 2)
    add_para(doc, "The sequence diagram shows the process of creating a reservation. The customer enters booking details through the web or mobile interface. The interface submits the request to the Laravel backend, which validates the input, checks the booking rules, stores the reservation, and returns a booking reference. The booking is then displayed as pending until the administrator reviews and updates the status.")
    add_image(doc, SEQUENCE_DIAGRAM, 6.4, "Sequence Diagram", "Sequence diagram showing customer reservation creation from the interface to Laravel API, database, and admin review.", "Figure 5.3. Sequence Diagram for Creating a Reservation")
    add_para(doc, "The sequence design confirms that the backend is responsible for validation and business logic. This is important because both the web and mobile applications depend on the same server-side rules when creating or updating reservations.")

    add_heading(doc, "5.3 System Implementation", 1)
    add_para(doc, "The system implementation resulted in a working Laravel web application and an Expo React Native mobile application. The Laravel backend handles authentication, routing, controllers, models, validation, reservation processing, payment proof upload, and database transactions. The Vue.js and Inertia.js web interface provides customer pages, booking forms, dashboards, staff pages, and administrator pages. The mobile application uses API endpoints to access the same reservation data and business rules.")
    add_para(doc, "The implemented customer module allows users to create an account, sign in, select event details, choose a branch and schedule, select a package, add menu items and add-ons, upload payment proof, submit the booking, view dashboard statistics, reschedule bookings, cancel bookings, and access booking pass information. These features address the need for a more convenient and traceable reservation process.")
    add_para(doc, "The implemented staff module supports operational tasks such as viewing today's bookings, checking in guests through a booking code, updating service status, and recording service adjustments. This helps staff members coordinate event preparation and update operational records without relying only on manual notes.")
    add_para(doc, "The implemented administrator module provides booking review, status updates, crew assignment, confirmed event tracking, branch management, catalog management, availability monitoring, report viewing, timeline tracking, and account management. The admin tools make it easier to maintain booking rules and monitor reservation activities across branches.")
    add_table(
        doc,
        "Implemented System Modules",
        ["Module", "Implemented Features", "Discussion"],
        [
            ("Customer Module", "Registration, login, booking, payment proof upload, dashboard, reschedule, cancellation, booking pass.", "Supports self-service reservation and booking monitoring."),
            ("Staff Module", "Daily bookings, check-in, service status, service adjustments, preparation list.", "Supports branch-level event operations."),
            ("Admin Module", "Booking review, status update, crew assignment, catalog, branches, reports, availability, accounts.", "Supports management and monitoring of reservation activities."),
            ("API Layer", "Mobile login, booking options, dashboard, operations, profile, admin and staff actions.", "Connects the Android app to the Laravel backend."),
            ("Database Layer", "Users, reservations, branches, event types, room options, menu items, add-ons, settings.", "Provides persistent storage for system records."),
        ],
    )

    add_heading(doc, "5.4 Evaluation of the System", 1)
    add_para(doc, "The system was evaluated based on selected ISO/IEC 25010 software quality criteria, particularly functional suitability, usability, performance efficiency, reliability, and security. Internal testing was conducted by checking whether the major system functions worked according to the requirements defined in Chapter 4. The evaluation also considered whether users could understand the interface flow and whether important actions returned appropriate feedback messages.")
    add_para(doc, "In terms of functional suitability, the system was able to perform the main reservation-related functions such as account access, booking creation, payment proof upload, reservation review, staff check-in, and status updating. The use of shared backend logic for web and mobile helped maintain consistent behavior across platforms.")
    add_para(doc, "In terms of usability, the system used guided forms, card-based layouts, status badges, tab navigation, and clear action buttons to help users complete tasks. The mobile application was especially designed around short task flows, such as booking, dashboard checking, account access, and staff operations.")
    add_para(doc, "In terms of performance efficiency, the system used API requests, cached mobile data, and structured database queries to reduce repeated loading. The mobile dashboard and booking options include caching behavior so that users can still see recently loaded information while the app refreshes data from the server.")
    add_para(doc, "In terms of reliability and security, the system used validation, Laravel authentication, Laravel Sanctum tokens for the mobile API, protected routes, role-based access, and error handling. These controls help prevent unauthorized access and reduce failures caused by incomplete or invalid input.")
    add_table(
        doc,
        "Evaluation Summary",
        ["Quality Criterion", "Observed Result", "Discussion"],
        [
            ("Functional Suitability", "Major customer, staff, and admin functions were implemented.", "The system supports the reservation workflow from booking creation to admin review and staff operation."),
            ("Usability", "Interface uses guided cards, labels, buttons, tabs, and status indicators.", "The design helps users understand the current task and booking status."),
            ("Performance Efficiency", "Mobile screens use API calls and caching for dashboard and booking data.", "The app can reduce repeated loading while still refreshing live records."),
            ("Reliability", "Validation and error messages are included in booking, login, and update actions.", "The system can guide users when required fields or network requests fail."),
            ("Security", "Authentication, role separation, and token-based mobile access are used.", "Private customer, staff, and admin functions are protected."),
        ],
    )
    add_para(doc, "The evaluation indicates that the system is appropriate for supporting digital reservation management. However, future testing with a larger number of respondents is recommended to obtain formal usability scores, task completion rates, and performance measurements under real operating conditions.")

    add_heading(doc, "5.5 Android Mobile Application", 1)
    add_para(doc, "The Android mobile application was developed using Expo, React Native, and TypeScript. Its current design follows the McDonald's-inspired visual identity used in the project, including a yellow customer page background, a header with the McLogo, cream and white rounded cards, red primary action buttons, chip selections, metric tiles, and bottom tab navigation. The app is designed for customers who need quick booking access and for staff or administrators who need mobile operational tools.")
    add_para(doc, "The customer-facing mobile flow includes Home, Book, Dashboard, and Account tabs. The Home screen introduces booking options and entry points. The Book screen guides users through event type selection, branch selection, schedule and availability checking, room choice, package selection, menu bundle selection, manual menu item selection, add-on selection, notes, and payment proof upload. The Dashboard screen displays booking metrics and reservation cards, while the Account screen supports profile and session management.")
    add_para(doc, "The mobile booking screen reflects the actual system process. It loads booking options from the Laravel API, stores cached booking data, calculates available start times based on duration, allows customers to choose menu items and add-ons, and submits a reservation using FormData so that payment proof images can be uploaded. This makes the mobile app a practical extension of the web reservation system.")
    add_para(doc, "The mobile dashboard also uses API-based data loading and caching. It displays upcoming reservations, confirmed spend, pending approvals, booking reference numbers, branch information, event dates, event times, and booking status labels. Customers can refresh the dashboard, add a new booking, cancel a reservation, or submit reschedule details when allowed.")
    add_para(doc, "The operations screen provides mobile access for staff and administrators. Staff can check in guests and update service status, while administrators can update booking status, assign crew, manage users, branches, inventory, event types, packages, room options, and booking settings. This design allows selected operational and administrative functions to be accessed even outside the desktop web dashboard.")
    create_mobile_ui(MOBILE_UI)
    add_image(doc, MOBILE_UI, 6.4, "Android Mobile Application Design", "Current mobile app design showing Home, Book, Dashboard, and Account screens using yellow customer pages, McLogo headers, rounded cards, chips, red buttons, and bottom tabs.", "Figure 5.4. Android Mobile Application Design")
    add_para(doc, "Overall, the Android mobile application strengthens the proposed system by making the reservation workflow more accessible. Customers can create and monitor bookings from a phone, staff can handle check-in and service updates, and administrators can perform selected management actions through mobile API screens. This supports the project's goal of making the reservation process more organized, traceable, and convenient.")

    doc.save(str(OUT))


if __name__ == "__main__":
    build_doc()
    print(OUT)
