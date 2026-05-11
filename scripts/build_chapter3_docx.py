from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
TEMPLATE = Path(r"C:\Users\User\Downloads\BSIT Capstone Project Template.docx")
OUT = ROOT / "Chapter 3 - Technical Background.docx"
DIAGRAM = ROOT / "scripts" / "chapter3_conceptual_framework.png"


ACCENT = RGBColor(191, 31, 36)
GOLD = RGBColor(255, 199, 44)
DARK = RGBColor(31, 31, 31)
LIGHT = "FFF7E6"
HEADER = "F4E1C1"


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = DARK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_para(doc, text: str, style: str = "Normal", align=None, bold: bool = False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5) if style == "Normal" else None
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = bold
    run.font.color.rgb = DARK
    return paragraph


def add_heading(doc, text: str, level: int):
    style = "Heading {}".format(level)
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    run.font.color.rgb = ACCENT if level == 1 else DARK
    run.font.size = Pt(14 if level == 1 else 12)
    return paragraph


def add_bullet(doc, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11.5)


def create_diagram(path: Path) -> None:
    if Image is None:
        return

    width, height = 1500, 850
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    try:
        title_font = ImageFont.truetype("arialbd.ttf", 38)
        box_font = ImageFont.truetype("arialbd.ttf", 25)
        small_font = ImageFont.truetype("arial.ttf", 22)
        caption_font = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        title_font = box_font = small_font = caption_font = ImageFont.load_default()

    draw.rounded_rectangle((28, 28, width - 28, height - 28), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "Conceptual Framework of the McDonald's Reservation System", fill=(31, 31, 31), font=title_font)

    boxes = {
        "customer": (70, 180, 400, 340, "Customer Module", "Booking, payment proof,\ndashboard, QR pass"),
        "staff": (70, 470, 400, 630, "Staff Module", "Check-in, prep list,\nservice updates"),
        "web": (560, 160, 940, 300, "Web Application Layer", "Laravel + Inertia + Vue\nresponsive pages"),
        "mobile": (560, 365, 940, 505, "Mobile Application Layer", "Expo + React Native\nAPI-driven app"),
        "api": (560, 595, 940, 735, "API and Business Logic", "Controllers, validation,\npricing, availability"),
        "admin": (1100, 180, 1430, 340, "Admin Module", "Bookings, catalog,\nbranches, reports"),
        "db": (1100, 470, 1430, 630, "Database and Storage", "Users, reservations,\nbranches, catalog,\npayment files"),
    }

    def box(coords, title, body, fill):
        draw.rounded_rectangle(coords, radius=18, fill=fill, outline=(184, 184, 184), width=3)
        x1, y1, x2, _ = coords
        draw.text((x1 + 24, y1 + 22), title, fill=(31, 31, 31), font=box_font)
        draw.multiline_text((x1 + 24, y1 + 66), body, fill=(55, 55, 55), font=small_font, spacing=5)

    for key, coords in boxes.items():
        fill = (255, 247, 230) if key in ("customer", "staff", "admin") else (255, 255, 255)
        if key == "api":
            fill = (255, 244, 204)
        box(coords[:4], coords[4], coords[5], fill)

    def arrow(start, end):
        draw.line((start, end), fill=(191, 31, 36), width=5)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            points = [(ex, ey), (ex - 18, ey - 11), (ex - 18, ey + 11)]
        else:
            points = [(ex, ey), (ex + 18, ey - 11), (ex + 18, ey + 11)]
        draw.polygon(points, fill=(191, 31, 36))

    arrow((400, 260), (560, 230))
    arrow((400, 550), (560, 435))
    arrow((750, 300), (750, 365))
    arrow((750, 505), (750, 595))
    arrow((940, 665), (1100, 550))
    arrow((940, 230), (1100, 260))
    arrow((1270, 470), (1270, 340))
    arrow((1100, 550), (940, 665))
    draw.text((60, 775), "Figure 3.1. Conceptual framework showing how customer, staff, and admin modules connect through the application, API, and database layers.", fill=(80, 80, 80), font=caption_font)
    image.save(path, quality=95)


def build_doc() -> None:
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("3. TECHNICAL BACKGROUND")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = ACCENT

    add_para(
        doc,
        "The proposed McDonald's Reservation System is a web and mobile-based reservation management system designed to support customers, branch staff, and administrators in handling event bookings. The technical foundation of the project is centered on a Laravel backend, an Inertia and Vue.js web interface, and an Expo React Native mobile application. These technologies work together to provide online booking, branch availability checking, payment proof uploading, booking approval, staff check-in, service monitoring, and customer account management.",
    )
    add_para(
        doc,
        "The system follows a client-server architecture. The client side is composed of the web application and the mobile application, while the server side is responsible for authentication, reservation processing, catalog management, business rules, and database transactions. This structure allows the system to separate the user interface from the application logic, making it easier to maintain, test, and improve the project as additional reservation features are introduced.",
    )
    add_para(
        doc,
        "For the web platform, the project uses PHP with the Laravel framework. Laravel provides routing, controllers, models, validation, authentication support, file storage, and database access through Eloquent ORM. These features are important because the reservation system needs to manage different types of data such as users, reservations, branches, packages, rooms, menu bundles, add-ons, payment proof files, and staff assignments. Laravel also supports role-based workflows, allowing customer, staff, and admin features to be separated while still sharing the same backend.",
    )
    add_para(
        doc,
        "The web user interface is developed using Vue.js with Inertia.js. Vue.js is used to build reactive pages such as the booking wizard, customer dashboard, admin booking review, staff dashboard, availability calendar, and profile management screens. Inertia.js connects Laravel and Vue without requiring a separate single-page application backend, allowing the system to behave like a modern application while still using Laravel routes and controllers. Tailwind CSS is used for styling because it provides reusable utility classes for creating responsive layouts and consistent interface components.",
    )
    add_para(
        doc,
        "The mobile application is developed using Expo and React Native. Expo provides the development environment for building the mobile app, while React Native allows the project to create native-style screens using TypeScript and reusable components. The mobile app communicates with the Laravel backend through API routes. It includes customer booking functions, dashboard viewing, profile management, staff operations, and admin actions such as updating booking status, assigning crew, and managing branch or catalog information.",
    )
    add_para(
        doc,
        "The system uses Laravel Sanctum for API authentication. When a mobile user logs in or registers, the backend issues a secure token that is included in succeeding API requests. This token-based approach protects private functions such as reservation creation, cancellation, rescheduling, staff check-in, profile updates, and admin management. On the mobile side, AsyncStorage is used to store the session token and cached user data so that the app can restore a user session and reduce repeated loading.",
    )
    add_para(
        doc,
        "The database layer stores the main records needed by the system. Important entities include User, Reservation, Branch, BookingPackage, EventType, RoomOption, MenuBundle, MenuItem, AddOn, and BookingSetting. The Reservation model contains booking details such as event date, time, duration, branch, number of guests, selected package, manual menu items, add-ons, payment proof path, total amount, booking reference, check-in code, booking status, and service status. This structure supports both customer-facing booking and operations-facing monitoring.",
    )
    add_para(
        doc,
        "Several technical concepts are applied in the project. First, role-based access control is used to separate customer, staff, and administrator capabilities. Second, API-based communication allows the mobile app to reuse the same backend data and business rules used by the web application. Third, reservation availability logic helps prevent scheduling conflicts by checking branch, date, time, duration, and capacity limits. Fourth, file upload handling is used for payment proof submission. Lastly, QR or booking pass functionality supports faster event check-in by giving staff a code-based way to verify arriving customers.",
    )

    add_heading(doc, "Technical Tools and Frameworks Used", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    headers = ["Component", "Technology", "Purpose in the System"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], HEADER)
    mark_header_row(table.rows[0])

    rows = [
        ("Backend Framework", "PHP and Laravel 9", "Handles routing, controllers, validation, models, authentication, business rules, and server-side reservation processing."),
        ("Web Frontend", "Vue.js, Inertia.js, Vite", "Provides responsive customer, staff, and admin web pages without separating the frontend from the Laravel application."),
        ("Mobile Frontend", "Expo, React Native, TypeScript", "Provides the mobile application for booking, dashboard viewing, staff operations, and selected admin functions."),
        ("Styling", "Tailwind CSS and reusable UI components", "Creates consistent layouts, forms, dashboards, status badges, and responsive visual design."),
        ("API Layer", "Laravel API routes and JSON responses", "Connects the mobile application to the backend for login, booking, profile, staff, and admin functions."),
        ("Authentication", "Laravel Sanctum and AsyncStorage", "Secures API requests through tokens and stores the mobile session locally on the device."),
        ("Database Layer", "Relational database through Eloquent ORM", "Stores users, reservations, branches, packages, menu items, add-ons, settings, and operational records."),
        ("File Storage", "Laravel storage for payment proof files", "Stores uploaded proof of payment and links it to the reservation record for admin review."),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT)

    add_heading(doc, "3.1 Conceptual Framework", 1)
    add_para(
        doc,
        "The conceptual framework of the proposed system presents the relationship among the main users, application modules, backend services, and database layer. The model shows that customers, staff, and administrators interact with the system through different interfaces, but their actions are processed by a shared application layer. This design reduces duplication because the booking rules, authentication process, reservation records, and catalog data are managed in one backend environment.",
    )

    create_diagram(DIAGRAM)
    if DIAGRAM.exists():
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = image_paragraph.add_run()
        inline_shape = run.add_picture(str(DIAGRAM), width=Inches(6.4))
        doc_pr = inline_shape._inline.docPr
        doc_pr.set("title", "Conceptual Framework Diagram")
        doc_pr.set(
            "descr",
            "Diagram showing customer, staff, and admin modules connecting to web and mobile application layers, API business logic, and database storage.",
        )
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        run = caption.add_run("Figure 3.1. Conceptual Framework of the McDonald's Reservation System")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.italic = True

    add_para(
        doc,
        "The Customer Module allows users to create an account, log in, select an event type, choose a branch, set the event date and time, select packages, add menu items, upload payment proof, and monitor booking status through the dashboard. This module directly addresses the problem of manual or inconvenient reservation processing by giving customers a guided booking workflow that can be accessed online or through the mobile application.",
    )
    add_para(
        doc,
        "The Staff Module supports daily branch operations. Staff users can view upcoming events, verify guests through a booking or check-in code, update the service status of an event, and record service adjustments when additional food, services, or time changes are needed. This module helps the branch handle confirmed events more efficiently because operational information is already connected to the reservation record.",
    )
    add_para(
        doc,
        "The Admin Module serves as the control center of the system. Administrators can review pending bookings, confirm or reject reservations, assign staff, manage branches, update event types, maintain booking packages, monitor availability, and generate operational reports. Through this module, management can maintain consistent reservation rules and track the overall performance of the booking process.",
    )
    add_para(
        doc,
        "The Web Application Layer and Mobile Application Layer are the access points of the system. The web layer is suitable for full dashboard management and administrative work, while the mobile layer provides convenient access for customers and staff who need to complete tasks quickly. Both layers communicate with the backend through Laravel routes and API endpoints, which process requests and return the needed data.",
    )
    add_para(
        doc,
        "The API and Business Logic Layer validates user input, applies booking rules, calculates totals, checks availability, stores uploaded payment proof, protects authenticated routes, and updates reservation status. This layer is important because it ensures that all user interfaces follow the same rules. For example, a reservation created from the mobile app and a reservation created from the web interface are processed using the same backend logic.",
    )
    add_para(
        doc,
        "The Database and Storage Layer keeps the persistent records of the system. It stores user accounts, roles, branch details, room choices, packages, menu selections, add-ons, reservations, payment proof locations, booking references, and check-in codes. Since all modules read from and write to this layer, it becomes the central source of accurate information for customers, staff, and administrators.",
    )
    add_para(
        doc,
        "Overall, the conceptual framework shows that the proposed system transforms the reservation process into an integrated digital workflow. Customers submit reservation requests, the backend validates and stores the data, administrators review and manage the booking, staff prepare and check in the event, and the database keeps the final record. This flow supports the goal of making McDonald's event reservations more organized, accessible, and easier to monitor.",
    )

    doc.save(str(OUT))


if __name__ == "__main__":
    build_doc()
    print(OUT)
