from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
TEMPLATE = Path(r"C:\Users\User\Downloads\BSIT Capstone Project Template.docx")
OUT = ROOT / "Chapter 6 - Conclusion and Recommendation.docx"
BULLET_NUM_ID = None


ACCENT = RGBColor(191, 31, 36)
DARK = RGBColor(31, 31, 31)


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(tag_name)):
        value = element.get(qn(attr_name))
        if value and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def ensure_bullet_numbering(doc) -> int:
    global BULLET_NUM_ID
    if BULLET_NUM_ID is not None:
        return BULLET_NUM_ID

    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, value in (("w:start", "1"), ("w:numFmt", "bullet"), ("w:lvlText", "\u2022"), ("w:lvlJc", "left")):
        element = OxmlElement(tag)
        element.set(qn("w:val"), value)
        lvl.append(element)

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    BULLET_NUM_ID = num_id
    return num_id


def add_para(doc, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return paragraph


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = ACCENT if level == 1 else DARK
    return paragraph


def add_bullet(doc, text: str):
    num_id = ensure_bullet_numbering(doc)
    paragraph = doc.add_paragraph(style="Normal")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK
    return paragraph


def build_doc():
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("6. CONCLUSION AND RECOMMENDATION")
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = ACCENT

    add_heading(doc, "6.1 Conclusion", 1)
    add_para(
        doc,
        "This study focused on the design and development of the McDonald's Event Planner, a digital reservation system intended to improve the process of booking birthday parties, business meetings, and other event reservations. The system was developed to address common problems in manual reservation handling, such as scheduling conflicts, incomplete booking records, delayed confirmation, limited monitoring, and difficulty coordinating customer requests with branch operations.",
    )
    add_para(
        doc,
        "Based on the results of the development process, the objectives of the study were met. The project produced a web-based reservation management system and an Android mobile application that support customer, staff, and administrator users. Customers can create accounts, submit reservation requests, select event details, choose packages and menu items, upload payment proof, monitor booking status, reschedule, cancel, and access booking pass information. Staff users can view event-related operational data, check in guests, and update service status. Administrators can review bookings, update reservation status, assign crew, manage branches, manage catalog information, monitor availability, and maintain user accounts.",
    )
    add_para(
        doc,
        "The system also achieved the objective of applying appropriate technologies to create a functional reservation platform. Laravel was used for the backend, database management, authentication, and business logic. Vue.js with Inertia.js was used for the web interface, while Expo React Native was used for the Android mobile application. Laravel Sanctum provided token-based API authentication for the mobile app, allowing the web and mobile platforms to share the same reservation data and backend rules.",
    )
    add_para(
        doc,
        "The study showed that a centralized reservation system can improve record management and operational coordination. By storing reservation details, customer information, branch schedules, payment proof, booking status, and service updates in one system, the proposed project reduces the risk of duplicate records and unclear communication. It also gives administrators and staff a clearer basis for preparing and monitoring events.",
    )
    add_para(
        doc,
        "The Android mobile application further supports the purpose of the study by making the reservation process more accessible to users. Its current design uses a McDonald's-inspired interface with yellow customer pages, rounded cards, booking chips, red action buttons, dashboard metrics, and bottom tab navigation. This makes the booking workflow easier to access through a mobile device and supports the increasing need for convenient digital service platforms.",
    )
    add_para(
        doc,
        "Overall, the proposed McDonald's Event Planner provides a practical solution for improving event reservation management. The system demonstrates how automation, centralized data storage, mobile access, and role-based features can help make the reservation process more organized, efficient, and user-friendly. Although the system can still be improved in future versions, the completed output shows that the study successfully addressed the main reservation problems identified in the project.",
    )

    add_heading(doc, "6.2 Recommendation", 1)
    add_para(
        doc,
        "Based on the findings and completed system output, several recommendations are proposed for future improvement. These recommendations aim to strengthen the system's usability, compatibility, security, scalability, and operational value. They may also guide future researchers who will continue or improve the proposed reservation platform.",
    )
    add_para(doc, "The following improvements are recommended:")
    recommendations = [
        "Develop an iOS version of the mobile application so that customers using iPhones can access the same booking and dashboard features.",
        "Add real-time notification features through email, SMS, or push notifications to inform customers about booking approval, rejection, rescheduling, cancellation, and payment proof status.",
        "Integrate online payment gateways so customers can pay reservation fees directly through the system instead of uploading payment proof manually.",
        "Improve the availability calendar by adding stronger conflict detection, branch capacity visualization, and clearer indicators for fully booked, available, and partially available schedules.",
        "Add QR code generation and scanning improvements for faster guest check-in during the event day.",
        "Include more detailed reporting and analytics, such as monthly reservation count, most selected package, branch performance, cancelled bookings, and customer demand patterns.",
        "Enhance security by adding stronger account verification, audit logs for administrator actions, and optional two-factor authentication for staff and admin accounts.",
        "Conduct formal usability testing with a larger group of respondents, including customers, branch staff, and administrators, to obtain measurable evaluation results.",
        "Improve offline or low-connectivity handling for the mobile application so recently loaded dashboard and booking information remains accessible when the network connection is unstable.",
        "Expand the system to support multiple branches more deeply, including branch-specific packages, staff schedules, room availability, and inventory preparation.",
    ]
    for item in recommendations:
        add_bullet(doc, item)
    add_para(
        doc,
        "Future researchers may also explore integration with customer loyalty programs, digital receipts, automated reminders, and AI-assisted package recommendations. These features can make the system more useful for both customers and management by personalizing the reservation experience and improving decision-making.",
    )
    add_para(
        doc,
        "It is also recommended that the system be tested in an actual branch environment before full deployment. Real-world testing can help identify issues that may not appear during development, such as network limitations, staff workflow differences, customer input errors, and branch-specific booking rules. Feedback from real users should be used to refine the system interface and improve the overall reservation process.",
    )
    add_para(
        doc,
        "Lastly, future development should continue following an iterative approach. Since customer expectations and business requirements may change over time, the system should be maintained and updated regularly. Continuous improvement will help ensure that the McDonald's Event Planner remains reliable, accessible, and effective for managing event reservations.",
    )

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    build_doc()
