from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\User\mcdonalds-reservation\docs\E-Learning Oral Defense Q&A.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color="B7C1CC", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_question_block(doc, number, question, answer):
    heading = doc.add_paragraph()
    heading.style = doc.styles["Heading 2"]
    run = heading.add_run(f"{number}. {question}")
    run.bold = True

    answer_p = doc.add_paragraph()
    answer_p.style = doc.styles["Normal"]
    answer_p.add_run("Suggested answer: ").bold = True
    answer_p.add_run(answer)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)
styles["Normal"].font.color.rgb = RGBColor(31, 41, 55)

styles["Title"].font.name = "Arial"
styles["Title"].font.size = Pt(22)
styles["Title"].font.bold = True
styles["Title"].font.color.rgb = RGBColor(17, 24, 39)

styles["Subtitle"].font.name = "Arial"
styles["Subtitle"].font.size = Pt(12)
styles["Subtitle"].font.color.rgb = RGBColor(75, 85, 99)

styles["Heading 1"].font.name = "Arial"
styles["Heading 1"].font.size = Pt(16)
styles["Heading 1"].font.bold = True
styles["Heading 1"].font.color.rgb = RGBColor(17, 24, 39)

styles["Heading 2"].font.name = "Arial"
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].font.bold = True
styles["Heading 2"].font.color.rgb = RGBColor(31, 41, 55)

header = section.header.paragraphs[0]
header.text = "E-Learning Oral Defense Q&A"
header.style = styles["Normal"]
header.runs[0].font.size = Pt(9)
header.runs[0].font.color.rgb = RGBColor(107, 114, 128)
add_bottom_border(header)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("Page ")
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(107, 114, 128)
fld_char_1 = OxmlElement("w:fldChar")
fld_char_1.set(qn("w:fldCharType"), "begin")
instr_text = OxmlElement("w:instrText")
instr_text.text = "PAGE"
fld_char_2 = OxmlElement("w:fldChar")
fld_char_2.set(qn("w:fldCharType"), "end")
footer_run._r.append(fld_char_1)
footer_run._r.append(instr_text)
footer_run._r.append(fld_char_2)

title = doc.add_paragraph(style="Title")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Oral Defense Q&A")

subtitle = doc.add_paragraph(style="Subtitle")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Learning Against the Odds: How E-Learning Survived in the Philippines During the Pandemic")

authors = doc.add_paragraph(style="Subtitle")
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
authors.add_run("Prepared for Mariz Mae Pareja, Josephine Vidal, and Jay Christian Estampa")

doc.add_paragraph()

meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = "Table Grid"
meta_table.autofit = False
for row in meta_table.rows:
    set_cell_width(row.cells[0], 2300)
    set_cell_width(row.cells[1], 7060)
for cell in meta_table.rows[0].cells:
    set_cell_shading(cell, "EEF2F7")
meta_table.cell(0, 0).text = "Editorial Type"
meta_table.cell(0, 1).text = "Commendatory editorial"
meta_table.cell(1, 0).text = "Main Stand"
meta_table.cell(1, 1).text = (
    "E-learning was imperfect, but it kept Philippine education alive during the pandemic through flexibility, teacher effort, family support, and student resilience."
)
for row in meta_table.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10)

doc.add_paragraph()

doc.add_heading("Short Opening Statement", level=1)
opening = (
    "Good day. Our editorial argues that e-learning in the Philippines during the pandemic was not perfect, "
    "but it was still a meaningful success because it kept education going during a national crisis. We emphasize "
    "that its strength did not come from technology alone, but from adaptability, teacher creativity, family support, "
    "and student resilience. The system exposed inequalities, but it also showed the need for better infrastructure "
    "and reforms. Overall, e-learning proved that learning can continue beyond the traditional classroom."
)
doc.add_paragraph(opening)

doc.add_heading("Possible Oral Defense Questions", level=1)

qa = [
    (
        "What is the main argument of your editorial?",
        "Our main argument is that e-learning in the Philippines survived during the pandemic because of adaptability and cooperation. Even with poor internet access, lack of gadgets, and unprepared households, education continued through online classes, printed modules, radio, and TV-based instruction.",
    ),
    (
        "Why did you choose this topic?",
        "We chose this topic because education during the pandemic affected almost every Filipino student, teacher, and family. We wanted to highlight not only the struggles but also the resilience and creativity shown during that period.",
    ),
    (
        "Why is your editorial commendatory?",
        "It is commendatory because it recognizes and praises the efforts made to continue education despite difficult conditions. We are not saying the system was perfect; we are appreciating the determination of teachers, students, families, and institutions.",
    ),
    (
        "What does your title mean by 'Learning Against the Odds'?",
        "It means that learning continued even when the situation was very difficult. Students faced signal problems, lack of devices, and limited support, but education still found ways to continue.",
    ),
    (
        "Do you believe e-learning was successful in the Philippines?",
        "Yes, but in a limited and realistic sense. It was successful because it kept education alive during a crisis. However, it was not equally successful for everyone because many students still experienced barriers like poor connectivity and lack of resources.",
    ),
    (
        "What was the biggest challenge of e-learning?",
        "The biggest challenge was inequality in access. Students with stable internet and gadgets had more advantages, while students in rural or low-income areas often struggled to participate fully.",
    ),
    (
        "If e-learning exposed inequality, why do you still defend it?",
        "We defend it because the inequality was already present before the pandemic. E-learning did not create all those problems; it revealed them more clearly. Because of that, it pushed people to recognize the need for better infrastructure, resources, and education policies.",
    ),
    (
        "What role did teachers play in the success of e-learning?",
        "Teachers played a major role because they had to adjust their methods, use online tools, prepare modules, communicate with students, and create new assessment strategies. Their flexibility helped students stay engaged despite the limitations.",
    ),
    (
        "What role did parents and families play?",
        "Parents and families became learning partners. Many helped explain lessons, manage schedules, and support students emotionally. Although many were unprepared, their involvement helped keep learning possible at home.",
    ),
    (
        "Why was modular learning important?",
        "Modular learning was important because not all students had internet access. Printed modules allowed learners without stable connectivity to continue studying, making education more accessible during lockdowns.",
    ),
    (
        "What is the strongest evidence you used?",
        "One strong evidence is the Department of Education's Learning Continuity Plan, which showed that the Philippines used flexible learning methods such as modules, online classes, radio, and TV instruction to respond to the crisis.",
    ),
    (
        "What is the weakness of your argument?",
        "A possible weakness is that we focus more on the positive side of e-learning. However, we also acknowledge its problems, especially unequal access, unprepared households, and the difficulty of independent learning.",
    ),
    (
        "What did students gain from e-learning?",
        "Students developed independence, time management, self-discipline, and critical thinking. Since teachers were not always physically present, learners had to take more responsibility for their own education.",
    ),
    (
        "What lesson should the education system learn from the pandemic?",
        "The education system should learn that flexibility is necessary. Schools need better digital infrastructure, stronger teacher training, accessible learning materials, and support systems for students and families.",
    ),
    (
        "If face-to-face classes are available again, should e-learning continue?",
        "Yes, but not as a total replacement. E-learning can support face-to-face education through blended learning, online resources, digital assessments, and flexible learning options for students who need them.",
    ),
    (
        "What makes your editorial relevant today?",
        "It is relevant because the pandemic showed that education systems must be prepared for disruptions. The lessons from e-learning can help improve future education, especially in emergencies.",
    ),
    (
        "What is your conclusion in simple terms?",
        "Our conclusion is that e-learning was imperfect but valuable. It helped education survive during the pandemic and showed the resilience, creativity, and determination of Filipino learners, teachers, and families.",
    ),
]

for index, (question, answer) in enumerate(qa, start=1):
    add_question_block(doc, index, question, answer)

doc.add_heading("Answering Strategy", level=1)
strategy = doc.add_paragraph()
strategy.add_run("If the panel challenges your claim, answer with this pattern: ").bold = True
strategy.add_run(
    "\"We acknowledge that weakness. However, our point is not that e-learning was perfect, but that it served its main purpose during the crisis, which was to keep education going.\""
)

doc.add_heading("Quick Closing Statement", level=1)
doc.add_paragraph(
    "In conclusion, our editorial praises e-learning not because it solved every educational problem, but because it helped Filipino education survive an extraordinary crisis. Its greatest lesson is that education becomes stronger when schools, teachers, students, families, and institutions adapt together."
)

doc.save(OUTPUT)
print(OUTPUT)
