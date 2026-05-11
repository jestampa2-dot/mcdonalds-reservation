from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = ImageDraw = ImageFont = None


ROOT = Path(r"C:\Users\User\mcdonalds-reservation")
TEMPLATE = Path(r"C:\Users\User\Downloads\BSIT Capstone Project Template.docx")
OUT = ROOT / "Chapter 4 - Methodology.docx"
DIAGRAM = ROOT / "scripts" / "chapter4_rad_methodology.png"


ACCENT = RGBColor(191, 31, 36)
DARK = RGBColor(31, 31, 31)
HEADER = "F4E1C1"
LIGHT = "FFF7E6"
BULLET_NUM_ID = None


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.bold = bold
    run.font.color.rgb = DARK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9D9D9")


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    values = []
    for element in numbering.findall(qn(tag_name)):
        value = element.get(qn(attr_name))
        if value and value.isdigit():
            values.append(int(value))
    return (max(values) + 1) if values else 1


def ensure_bullet_numbering(doc) -> int:
    global BULLET_NUM_ID
    if BULLET_NUM_ID is not None:
        return BULLET_NUM_ID

    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "singleLevel")
    abstract.append(multi_level)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")

    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)

    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Symbol")
    r_fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(r_fonts)

    for item in (start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr):
        lvl.append(item)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    BULLET_NUM_ID = num_id
    return num_id


def add_para(doc, text: str):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    return paragraph


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = ACCENT if level == 1 else DARK
    return paragraph


def add_bullet(doc, text: str):
    num_id = ensure_bullet_numbering(doc)
    paragraph = doc.add_paragraph(style="Normal")
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    paragraph.paragraph_format.left_indent = Inches(0.45)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK
    return paragraph


def create_diagram(path: Path) -> None:
    if Image is None:
        return

    width, height = 1500, 720
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    try:
        title_font = ImageFont.truetype("arialbd.ttf", 38)
        box_font = ImageFont.truetype("arialbd.ttf", 25)
        small_font = ImageFont.truetype("arial.ttf", 21)
        caption_font = ImageFont.truetype("arial.ttf", 20)
    except Exception:
        title_font = box_font = small_font = caption_font = ImageFont.load_default()

    draw.rounded_rectangle((28, 28, width - 28, height - 28), radius=24, outline=(210, 210, 210), width=3)
    draw.text((60, 55), "RAD Methodology Applied to the McDonald's Reservation System", fill=(31, 31, 31), font=title_font)

    boxes = [
        (70, 205, 340, 405, "1. Planning\nRequirements", "Identify users,\nprocess gaps,\nand system needs"),
        (390, 205, 660, 405, "2. Prototype", "Prepare Figma\nwireframes and\ninterface flow"),
        (710, 205, 980, 405, "3. Receive\nFeedback", "Review prototype,\nrecord comments,\nand refine scope"),
        (1030, 205, 1300, 405, "4. Finalize\nSoftware", "Develop Laravel,\nVue, and Expo\nsystem features"),
        (500, 450, 1000, 645, "5. Evaluation Method and Tools", "Assess usability,\nperformance,\nsecurity, and\nfunctional suitability"),
    ]

    for index, (x1, y1, x2, y2, title, body) in enumerate(boxes):
        fill = (255, 247, 230) if index < 4 else (255, 244, 204)
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=(184, 184, 184), width=3)
        draw.multiline_text((x1 + 22, y1 + 22), title, fill=(31, 31, 31), font=box_font, spacing=4)
        draw.multiline_text((x1 + 22, y1 + 95), body, fill=(55, 55, 55), font=small_font, spacing=5)

    def arrow(start, end):
        draw.line((start, end), fill=(191, 31, 36), width=5)
        ex, ey = end
        sx, sy = start
        if ex >= sx:
            points = [(ex, ey), (ex - 18, ey - 11), (ex - 18, ey + 11)]
        else:
            points = [(ex, ey), (ex + 18, ey - 11), (ex + 18, ey + 11)]
        draw.polygon(points, fill=(191, 31, 36))

    arrow((340, 305), (390, 305))
    arrow((660, 305), (710, 305))
    arrow((980, 305), (1030, 305))
    arrow((1165, 405), (1000, 535))
    arrow((500, 535), (205, 405))
    draw.text((60, 655), "Figure 4.1. RAD-based methodology flow used in developing and evaluating the proposed reservation system.", fill=(80, 80, 80), font=caption_font)
    image.save(path, quality=95)


def add_table(doc):
    add_heading(doc, "Summary of Methodology Activities", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    for idx, header in enumerate(["RAD Phase", "Main Activity", "Expected Output"]):
        set_cell_text(table.rows[0].cells[idx], header, True)
        set_cell_shading(table.rows[0].cells[idx], HEADER)
    mark_header_row(table.rows[0])

    rows = [
        ("Planning Requirements", "Analyze reservation problems, identify users, and list functional and non-functional requirements.", "Requirements list, user roles, process scope, and feature priorities."),
        ("Prototype", "Create low-fidelity and interface flow designs for customer, staff, and admin screens.", "Figma wireframes, proposed screen flow, and prototype navigation."),
        ("Receive Feedback", "Present the prototype to possible users and improve the design based on comments.", "Revised prototype, clarified features, and improved workflow."),
        ("Finalize Software", "Develop the Laravel web system, Expo mobile app, database, API routes, and authentication.", "Working reservation system with customer, staff, and admin modules."),
        ("Evaluation", "Test the system using selected ISO 25010 criteria, task observation, questionnaires, and error logs.", "Evaluation results, usability feedback, and recommendations."),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
            if idx == 0:
                set_cell_shading(cells[idx], LIGHT)


def build_doc() -> None:
    doc = Document(str(TEMPLATE))
    clear_body(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(12)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    title_run = title.add_run("4. METHODOLOGY")
    title_run.font.name = "Arial"
    title_run.font.size = Pt(14)
    title_run.bold = True
    title_run.font.color.rgb = ACCENT

    add_para(
        doc,
        "This chapter presents the methodology used in the design and development of the proposed McDonald's Reservation System. The project followed the Rapid Application Development (RAD) methodology because the system requires quick prototyping, repeated user feedback, and continuous improvement of interface and functional requirements. RAD is suitable for this study because the reservation process involves several user groups, including customers, staff, and administrators, whose needs must be reviewed through visible prototypes and working software iterations.",
    )
    add_para(
        doc,
        "Based on Systems Analysis and Design concepts, the project began by identifying the current reservation problems, defining the proposed system scope, analyzing user roles, designing the system flow, developing prototypes, implementing the software, and evaluating the output. The methodology focuses on converting the existing manual or semi-manual reservation process into a structured digital workflow where booking requests, payment proof, schedule checking, admin approval, and staff event handling can be managed in one system.",
    )
    add_para(
        doc,
        "The RAD model was selected because it supports active participation from stakeholders during development. Instead of completing the entire system before showing it to users, the researchers created prototype screens and revised the workflow according to feedback. This approach helped the team identify missing features, unclear interface labels, possible booking conflicts, and operational needs before finalizing the Laravel web application and Expo React Native mobile application.",
    )

    create_diagram(DIAGRAM)
    if DIAGRAM.exists():
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        inline_shape = image_paragraph.add_run().add_picture(str(DIAGRAM), width=Inches(6.4))
        doc_pr = inline_shape._inline.docPr
        doc_pr.set("title", "RAD Methodology Diagram")
        doc_pr.set(
            "descr",
            "Diagram showing the RAD process from planning requirements to prototype, feedback, software finalization, and evaluation.",
        )
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        run = caption.add_run("Figure 4.1. RAD Methodology Applied to the Proposed System")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.italic = True

    add_table(doc)

    add_heading(doc, "4.1. Planning Requirements", 1)
    add_para(
        doc,
        "In the planning requirements phase, the researchers identified the main problems in handling McDonald's event reservations. The study considered the needs of customers who want a convenient way to book events, staff members who need organized schedules and check-in records, and administrators who need a reliable dashboard for reviewing bookings, branches, packages, reports, and availability. The purpose of this phase was to define what the system should do and what limits should be observed during development.",
    )
    add_para(
        doc,
        "Data for requirements planning may be gathered through informal interviews, observation of the existing reservation process, review of current booking forms or procedures, and consultation with potential users. The gathered information is then translated into functional and non-functional requirements. Functional requirements describe the features the system must provide, while non-functional requirements describe quality attributes such as usability, security, reliability, accessibility, and performance.",
    )
    add_para(doc, "Key requirements included:")
    for item in [
        "Customer registration, login, profile management, and secure session handling.",
        "Online reservation creation with event type, branch, date, time, duration, room choice, guest count, package, menu items, add-ons, notes, and payment proof upload.",
        "Availability checking to help prevent booking conflicts based on branch, date, time, duration, and capacity.",
        "Customer dashboard for viewing booking status, reservation details, rescheduling, cancellation, and QR or booking pass access.",
        "Staff dashboard for daily event preparation, guest check-in, service status updates, and service adjustments.",
        "Admin dashboard for pending booking review, confirmation, rejection, crew assignment, branch management, catalog management, availability monitoring, reports, and user account management.",
        "Secure authentication and role-based access control for customer, staff, and administrator users.",
        "Database storage for users, reservations, branches, packages, menu selections, add-ons, payment proof paths, booking references, and check-in codes.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "At the end of this phase, the researchers defined the system boundary. The proposed system focuses on reservation management, booking monitoring, staff operations, and admin control. It does not aim to replace all McDonald's point-of-sale or enterprise systems. This limitation keeps the project manageable while still addressing the main reservation workflow identified in the study.",
    )

    add_heading(doc, "4.2. Prototype", 1)
    add_para(
        doc,
        "In the prototype phase, the researchers prepared a low-fidelity design to represent the proposed system before full software development. Figma was used to organize the interface flow, screen layout, and user navigation for the customer, staff, and admin modules. The prototype served as a visual guide for understanding how users would move from one screen to another, such as from login to booking, from booking confirmation to dashboard, and from admin review to confirmed events.",
    )
    add_para(
        doc,
        "The prototype was important because it helped the researchers evaluate the system design before coding. Through wireframes, the team could check whether the booking steps were understandable, whether the fields were arranged logically, and whether the admin and staff screens contained the needed information. This reduced the risk of developing features that users might find confusing or incomplete.",
    )
    add_para(doc, "Mock-up elements created:")
    for item in [
        "Landing or home screen showing booking entry points and event options.",
        "Login and registration screens for customer access.",
        "Customer booking wizard for event details, branch selection, schedule, package, menu, add-ons, payment proof, and review.",
        "Customer dashboard, booking details, reschedule option, cancellation option, and QR or booking pass screen.",
        "Staff dashboard with daily events, check-in input, preparation list, and service status update controls.",
        "Admin dashboard with booking statistics, pending booking review, branch management, availability calendar, catalog management, reports, timeline, and account management.",
        "Mobile-first screens for the Expo React Native application and responsive layouts for web users.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The prototype also supported SAD modeling because it made the proposed process visible. It helped connect the requirements gathered in the planning phase to concrete interface elements, data inputs, and expected outputs. The prototype therefore acted as a bridge between system analysis and actual system implementation.",
    )

    add_heading(doc, "4.3. Receive Feedback", 1)
    add_para(
        doc,
        "After preparing the prototype, the researchers presented the design to possible users and stakeholders for comments and suggestions. The feedback phase focused on checking whether the proposed workflow was understandable, complete, and practical for real reservation activities. Users were asked to review the booking steps, dashboard information, admin approval flow, staff check-in process, and general screen layout.",
    )
    add_para(doc, "Feedback gathered may include suggestions such as:")
    for item in [
        "Simplifying the booking wizard so customers can complete reservation details step by step.",
        "Adding clearer status labels for pending, confirmed, cancelled, completed, and rejected reservations.",
        "Providing an availability calendar to help administrators monitor booked and open dates.",
        "Adding payment proof upload so administrators can review reservation payments before confirmation.",
        "Including a QR or booking pass feature to speed up guest verification during event check-in.",
        "Improving staff screens by showing daily events, preparation lists, and service status controls.",
        "Making the interface mobile-friendly for customers and staff who may use phones during booking or operations.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The researchers revised the prototype and feature priorities based on the feedback. Changes were made to improve navigation, reduce confusing steps, and ensure that customer, staff, and admin users each had access to the information they needed. This phase reflects the RAD principle of iterative improvement, where the system design evolves through user comments before final implementation.",
    )

    add_heading(doc, "4.4. Finalize Software", 1)
    add_para(
        doc,
        "In the finalize software phase, the researchers implemented the system using the selected technologies identified in the technical background. The backend was developed using PHP and Laravel, while the web interface used Vue.js, Inertia.js, Vite, and Tailwind CSS. The mobile application was developed using Expo, React Native, and TypeScript. The system database was structured through Laravel migrations and Eloquent models to manage users, reservations, branches, packages, menu items, add-ons, booking settings, and other operational records.",
    )
    add_para(doc, "The final application was developed using:")
    for item in [
        "Laravel for routing, controllers, validation, authentication, database models, file upload handling, and business logic.",
        "Vue.js with Inertia.js for the responsive web interface, including customer, staff, and admin dashboards.",
        "Expo and React Native with TypeScript for the mobile application interface and API-driven mobile workflows.",
        "Laravel Sanctum for secure mobile API authentication using token-based access.",
        "Tailwind CSS for consistent layout, spacing, colors, buttons, forms, cards, and status indicators.",
        "Relational database tables and Eloquent ORM for storing system records and managing relationships.",
        "Figma for prototype design and interface planning before implementation.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The finalized system included customer booking, dashboard viewing, payment proof upload, booking reschedule and cancellation, QR or booking pass access, staff check-in, service status updates, admin booking review, crew assignment, branch management, catalog management, availability monitoring, reports, and account management. The mobile application reused the Laravel backend through API endpoints, allowing the project to maintain consistent data and rules between the web and mobile platforms.",
    )
    add_para(
        doc,
        "Internal testing was conducted during development to check form validation, login and registration, API communication, booking creation, admin approval, staff check-in, file upload behavior, and role-based access. Errors found during testing were corrected by improving validation rules, interface messages, database handling, and user navigation. This ensured that the final version was more stable and easier to use before evaluation.",
    )

    add_heading(doc, "4.5 Evaluation Method and Tools", 1)
    add_para(
        doc,
        "To assess the quality of the proposed McDonald's Reservation System, the researchers used selected criteria from the ISO/IEC 25010 software quality model. The evaluation focused on functional suitability, usability, performance efficiency, reliability, and security. These criteria were chosen because the system must correctly process reservations, remain understandable to users, respond within an acceptable time, protect account access, and support stable booking operations.",
    )
    add_para(doc, "Evaluation tools and techniques included:")
    for item in [
        "User acceptance testing with selected customers, staff, and administrator representatives.",
        "System Usability Scale (SUS) or usability questionnaire to measure ease of use, clarity, and user satisfaction.",
        "Task completion observation for activities such as account login, reservation creation, admin approval, and staff check-in.",
        "Test cases for functional requirements, including booking, cancellation, rescheduling, payment proof upload, and status updates.",
        "Error logs and debugging records to identify failed requests, validation issues, and unexpected system behavior.",
        "Performance observation by checking the response time of important actions such as loading booking options, submitting reservations, and retrieving dashboard data.",
        "Security checks for authenticated routes, token-based mobile access, role restrictions, and protected user functions.",
    ]:
        add_bullet(doc, item)
    add_para(
        doc,
        "The respondents for evaluation may include customers who will test the booking process, staff users who will test check-in and service monitoring, and administrators who will test booking review and management features. Their responses will be summarized to determine whether the system meets the expected level of usability and functional completeness. The results will also guide future improvements, especially in interface design, mobile responsiveness, notification features, and reporting.",
    )
    add_para(
        doc,
        "The evaluation results will be interpreted using descriptive measures such as frequency, mean score, task completion rate, and summary of user comments. If the system receives acceptable ratings in usability, reliability, performance efficiency, security, and functional suitability, it can be considered effective for supporting the reservation process. Any issues found during evaluation will be documented as recommendations for enhancement and future development.",
    )

    doc.save(str(OUT))


if __name__ == "__main__":
    build_doc()
    print(OUT)
