from pathlib import Path
from urllib.parse import quote_plus

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


SOURCE = Path(r"C:\Users\User\OneDrive\Documents\Cpastone Project\Chapter2.docx")
OUTPUT = Path(r"C:\Users\User\mcdonalds-reservation\Chapter2 - Hyperlinked References.docx")


VERIFIED = [
    (
        "Delizo, G. A., & Esguerra, M. A. (2013). Online hotel reservation and management system for the College of International Tourism and Hospitality Management (CITHM). International Journal of Computers & Technology, 10(1), 1201-1229.",
        "https://doi.org/10.24297/ijct.v10i1.3324",
        "Note: Your chapter cites this as 2011, but the online journal record shows 2013.",
    ),
    (
        "Felicen, S. S., & Vizconde, A. M. L. (2012). Computerized reservation system (CRS) as a feature of business operation of travel agencies.",
        "https://research.lpubatangas.edu.ph/wp-content/uploads/2014/08/JTHR-Computerized-Reservation-System.pdf",
        "",
    ),
    (
        "Ozturk, A. B., Bilgihan, A., Nusair, K., & Okumus, F. (2016). What keeps the mobile hotel booking users loyal? Investigating the roles of self-efficacy, compatibility, perceived ease of use, and perceived convenience.",
        "https://doi.org/10.1016/j.ijinfomgt.2016.04.005",
        "",
    ),
    (
        "Law, R., Buhalis, D., & Cobanoglu, C. (2014). Progress on information and communication technologies in hospitality and tourism.",
        "https://doi.org/10.1108/IJCHM-08-2013-0367",
        "Note: This is the closest verified match to the Law-related ICT hospitality citation in the chapter.",
    ),
    (
        "Ganiyu, M., Egwuche, O. S., & Adekunle, K. O. (2022/2024). Online events booking and reservation system.",
        "https://www.researchgate.net/publication/383272017_Online_Events_Booking_and_Reservation_System",
        "Note: ResearchGate shows the conference paper as August 2022, uploaded in 2024.",
    ),
    (
        "Egigogo, R. A., Tijjani Naniya, M., Ahmad Abubakar, A., & Mansir, A. (2024). Design and implementation of computerized restaurant table booking system. Ceddi Journal of Information System and Technology, 3(1), 47-55.",
        "https://doi.org/10.56134/jst.v3i1.64",
        "",
    ),
    (
        "Deksne, L., Kempelis, A., Sniedzins, T., & Kozlovskis, A. (2021). Automated system for restaurant services. Information Technology and Management Science, 24, 15-25.",
        "https://doi.org/10.7250/itms-2021-0003",
        "",
    ),
    (
        "Subrahmanyam, K., Viswanadham, D., Sujith, A., Devavarapu, R., & Venkat, K. S. (2021). Online table reservation with pre-ordering.",
        "https://www.proquest.com/docview/2623049632/C88C66C6B1814EE7PQ/3",
        "",
    ),
    (
        "Sahara, R. I., Chairunnisa, R., Muchamad Iqbal, M. I., & Azmi Hardi Roza, A. H. (2025). Development of a web-based reservation system to improve the efficiency of catering services. bit-Tech, 8(1), 799-808.",
        "https://doi.org/10.32877/bt.v8i1.2718",
        "",
    ),
    (
        "Davydova, O., & Sysoieva, S. (2025). Specifics of hotel and restaurant business management using modern information technologies.",
        "https://gst-journal.net.ua/index.php/gst/article/view/58",
        "",
    ),
    (
        "Rochella, M., & Tileng, K. G. (2026). Design and development of the Restaurant X reservation application on the iOS platform using App Clip. Sistemasi, 15(2), 669-683.",
        "https://doi.org/10.32520/stmsi.v15i2.5953",
        "",
    ),
    (
        "Luzzi, M., Guerriero, F., Macrina, G., & Triki, C. (2026). Restaurant revenue management through combinatorial auctions. Annals of Operations Research.",
        "https://doi.org/10.1007/s10479-026-07198-7",
        "",
    ),
    (
        "Horozhankina, N. A., Razinkova, M. Yu., Sabirov, O. V., & Mulenko, K. O. (2026). Innovative activities in the hotel and restaurant business.",
        "https://doi.org/10.32782/2521-6643-2026-1-71.26",
        "",
    ),
    (
        "Samanta, R. N., Mathad, V. P., & Kamath, R. (2026). Reception automation and guest experience: Balancing efficiency and personalization in hotel front offices.",
        "https://eurekamag.com/research/104/595/104595466.php",
        "",
    ),
    (
        "Gheorghe, N.-D., Alexandrescu, B., & Constantinescu, R.-C. (2025). A web-based innovative approach of the booking system for the food industry.",
        "https://eeae-conf.uni-ruse.bg/images/files/2025_VirtualPosters/VP31.pdf",
        "",
    ),
]


UNVERIFIED = [
    "Cruz and Santos (2015)",
    "Reyes and Lopez (2017)",
    "Mendoza (2018)",
    "Bautista and Garcia (2016)",
    "Torres (2019)",
    "Dela Cruz et al. (2020)",
    "Dublado et al. (2014)",
    "Fernandez et al. (2019)",
    "Esguerra and Arreza (2021)",
    "Alonzo et al. (2020)",
    "Ramos and Delgado (2021)",
    "Villarta (2020)",
    "Dizon and Morales (2019)",
    "Aquino and Reyes (2022)",
    "Torres and Mendoza (2021)",
    "Bautista (2020)",
    "Salonga and Cruz (2018)",
    "Perez and Navarro (2021)",
    "Herrera and Gomez (2022)",
    "Flores (2019)",
    "Robles and Evangelista (2021)",
    "Lim (2020)",
    "Cruz and Villanueva (2022)",
    "Zezar and Kalifa (2025)",
    "Meena and Rana (2026)",
    "Babu, Monika, Lakshmi, and Swathi (2024)",
    "Jun, Iqbal, and Subramaniam (2023)",
    "Gokul, Ishnu, and Logeshwara (2026)",
]


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_normal(paragraph, indent=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(18) if indent else Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(-18) if indent else Pt(0)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)


def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 31, 31)
    return paragraph


def build():
    doc = Document(str(SOURCE))

    doc.add_page_break()
    add_heading(doc, "References with Hyperlinks")

    note = doc.add_paragraph()
    note.add_run(
        "The following references were matched with credible online records during verification. Some citation years or author names in the literature may need correction based on the linked source records."
    )
    set_normal(note, indent=False)

    for reference, url, note_text in VERIFIED:
        p = doc.add_paragraph()
        p.add_run(reference + " ")
        add_hyperlink(p, "Open source", url)
        set_normal(p)
        if note_text:
            n = doc.add_paragraph()
            n.add_run(note_text)
            set_normal(n, indent=False)

    add_heading(doc, "Citations Needing Verification")
    note2 = doc.add_paragraph()
    note2.add_run(
        "These citations appear in the related literature, but I could not verify an exact credible online source from the available search results. The links below are search links only, not confirmed references."
    )
    set_normal(note2, indent=False)

    for citation in UNVERIFIED:
        query = f"{citation} reservation system hospitality restaurant"
        p = doc.add_paragraph()
        p.add_run(citation + " - ")
        add_hyperlink(p, "Search online", "https://scholar.google.com/scholar?q=" + quote_plus(query))
        set_normal(p)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    build()
