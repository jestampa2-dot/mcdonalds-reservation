from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


SOURCE = Path(r"C:\Users\User\OneDrive\Documents\Cpastone Project\Chapter2.docx")
OUTPUT = Path(r"C:\Users\User\mcdonalds-reservation\Chapter2 - Organized with Modules.docx")

MODULES_BEFORE_PARAGRAPH = {
    1: "Module 1: Overview of Related Literature",
    2: "Module 2: Summary of Local and Foreign Reservation Studies",
    7: "Module 3: Research Gap",
    9: "Module 4: Local Literature - Automation in Reservation and Hospitality Systems",
    11: "Module 5: Local Literature - Centralized Booking and Availability Management",
    14: "Module 6: Local Literature - Institutional, Hospitality, and Management Information Systems",
    16: "Module 7: Local Literature - User Experience, Web-Based Access, and Mobile Compatibility",
    27: "Module 8: Foreign Literature - Online Event Booking and Hospitality Platforms",
    29: "Module 9: Foreign Literature - Digital Transformation and ICT Integration",
    32: "Module 10: Foreign Literature - Web-Based Restaurant and Catering Reservation Systems",
    36: "Module 11: Foreign Literature - Mobile Access, Automation, Capacity, and Emerging Technologies",
}


def insert_paragraph_before(paragraph, text):
    new_element = OxmlElement("w:p")
    paragraph._p.addprevious(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    run = new_paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 31, 31)
    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    new_paragraph.paragraph_format.space_before = Pt(10)
    new_paragraph.paragraph_format.space_after = Pt(4)
    return new_paragraph


def organize():
    doc = Document(str(SOURCE))
    original_paragraphs = list(doc.paragraphs)

    for index in sorted(MODULES_BEFORE_PARAGRAPH.keys(), reverse=True):
        if index < len(original_paragraphs):
            insert_paragraph_before(original_paragraphs[index], MODULES_BEFORE_PARAGRAPH[index])

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    organize()
